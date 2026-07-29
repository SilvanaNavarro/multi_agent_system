import json
import os
import re
import subprocess
import threading
from dotenv import load_dotenv
load_dotenv()

# ============================================================
# MODOS DEL AGENTE
# ============================================================
MODOS = {
    "default":        "agent_prompt_default.txt",
    "ciberseguridad": "agent_prompt_ciberseguridad.txt",
    "fullstack":      "agent_prompt_fullstack.txt",
    "data_engineer":  "agent_prompt_data_engineer.txt",
    "devops":         "agent_prompt_devops.txt",
    "agente_creator": "agent_prompt_agente_creator.txt",
}

MODOS_ETIQUETAS = {
    "default":        "Project Manager",
    "ciberseguridad": "Experto en Ciberseguridad",
    "fullstack":      "Desarrollador Fullstack",
    "data_engineer":  "Ingeniero de Datos",
    "devops":         "Ingeniero DevOps",
    "agente_creator": "KIKE – Creador de Agentes",
}

MODO_ACTUAL = "default"
CUSTOM_MODES_FILE = "custom_modes.json"
AGENT_DIR = os.path.abspath(os.path.dirname(__file__))
RUTA_PROYECTO = [None]


def cargar_prompt(modo):
    archivo = MODOS.get(modo, MODOS["default"])
    with open(archivo, "r", encoding="utf-8") as f:
        return f.read().strip()


def cargar_conocimiento(modo):
    archivo = f"knowledge_{modo}.json"
    if not os.path.exists(archivo):
        return []
    with open(archivo, "r", encoding="utf-8") as f:
        return json.load(f).get("entries", [])


def guardar_conocimiento(modo, entries):
    archivo = f"knowledge_{modo}.json"
    with open(archivo, "w", encoding="utf-8") as f:
        json.dump({"entries": entries}, f, indent=2, ensure_ascii=False)


def construir_system_prompt(modo):
    base = cargar_prompt(modo)
    entries = cargar_conocimiento(modo)

    modos_lista = "\n".join(f'  - "{k}" → {v}' for k, v in MODOS_ETIQUETAS.items() if k != modo)
    switching_rule = (
        "\n\n## REGLA DE CAMBIO DE EXPERTO\n"
        "PRIORIDAD: Antes de cambiar de modo, verifica si alguna herramienta disponible puede resolver la solicitud.\n"
        "- Cambia de modo SOLO si ninguna herramienta existente puede manejar la solicitud Y el tema es claramente de otro experto.\n"
        "- Si existe una herramienta adecuada para la tarea, ÚSALA sin cambiar de modo.\n\n"
        "Si debes cambiar de modo, usa:\n"
        '<tool_call>{"name": "cambiar_modo", "inputs": {"modo": "nombre_modo"}}</tool_call>\n'
        f"Expertos disponibles:\n{modos_lista}"
    )

    approval_rule = (
        "\n\n## REGLA DE APROBACIÓN OBLIGATORIA — MODIFICACIÓN DE ARCHIVOS\n"
        "ANTES de llamar a editar_archivo o crear_archivo SIEMPRE debes:\n"
        "1. Escribir un mensaje al usuario explicando QUÉ cambios planeas hacer y POR QUÉ.\n"
        "2. Llamar a pedir_confirmacion con un resumen de TODOS los cambios planeados.\n"
        "3. Solo si el usuario aprueba, ejecutar las herramientas de edición.\n"
        "NUNCA edites ni crees archivos sin haber explicado el plan y recibido aprobación explícita.\n"
        "Esta regla es ABSOLUTA — aplica sin excepción a todos los modos y backends."
    )

    if not entries:
        return base + switching_rule + approval_rule
    conocimiento_txt = "\n".join(f"- [{e['tema']}]: {e['contenido']}" for e in entries)
    return f"{base}{switching_rule}{approval_rule}\n\n## CONOCIMIENTO ACUMULADO DE EXPERIENCIAS PREVIAS\n{conocimiento_txt}"


SYSTEM_PROMPT = construir_system_prompt(MODO_ACTUAL)

# Callbacks de la GUI — se setean en __main__ antes de arrancar
# _solicitar_permiso:     fn(nombre_modo, etiqueta, contenido_prompt) -> bool
# _solicitar_ruta:        fn() -> str | None
# _solicitar_confirmacion: fn(pregunta) -> bool
_solicitar_permiso      = [None]
_solicitar_ruta         = [None]
_solicitar_confirmacion = [None]


# ============================================================
# BACKEND
# ============================================================
BACKEND = "zhipu"  # "anthropic" | "ollama" | "claude-code" | "zhipu"
MODELO  = ""
client  = None

BACKENDS_DISPONIBLES = ["claude-code", "anthropic", "ollama", "zhipu"]

MODELOS_POR_BACKEND = {
    "anthropic":   "claude-opus-4-8",
    "claude-code": "claude-opus-4-8",
    "ollama":      "llama3.2",
    "zhipu":       "glm-4.5-flash",
}


def inicializar_cliente(backend):
    global client, MODELO, BACKEND
    BACKEND = backend
    MODELO  = MODELOS_POR_BACKEND.get(backend, "llama3.2")
    if backend == "anthropic":
        try:
            from anthropic import Anthropic
            client = Anthropic()
        except ImportError:
            return "Paquete 'anthropic' no instalado. Ejecuta: pip install anthropic"
    elif backend == "claude-code":
        client = None
    elif backend == "ollama":
        try:
            from ollama import Client as OllamaClient
            client = OllamaClient()
        except ImportError:
            return "Paquete 'ollama' no instalado. Ejecuta: pip install ollama"
    elif backend == "zhipu":
        try:
            from zai import ZaiClient
            api_key = os.environ.get("ZAI_API_KEY", "")
            if not api_key:
                return "Variable ZAI_API_KEY no configurada. Ejecuta: export ZAI_API_KEY=tu_clave"
            client = ZaiClient(api_key=api_key)
        except ImportError:
            return "Paquete 'zai-sdk' no instalado. Ejecuta: pip install zai-sdk"
    return None


_error_inicializacion = inicializar_cliente(BACKEND)
if _error_inicializacion:
    print(f"[Advertencia backend '{BACKEND}']: {_error_inicializacion}")


# ============================================================
# HERRAMIENTAS
# ============================================================
HERRAMIENTAS_DINAMICAS = {}
CUSTOM_TOOLS_FILE = "custom_tools.json"
TOOLS_DESACTIVADAS: set = set()  # nombres de herramientas desactivadas desde la GUI

TOOLS = [
    {
        "name": "calcular",
        "description": "Realiza operaciones matemáticas básicas",
        "input_schema": {
            "type": "object",
            "properties": {
                "operacion": {"type": "string", "description": "Expresión a calcular, ej: '2 + 2'"}
            },
            "required": ["operacion"]
        }
    },
    {
        "name": "cambiar_modo",
        "description": (
            "Cambia el modo del agente activando al experto del área requerida. "
            "IMPORTANTE: Usar SOLO cuando ninguna herramienta disponible puede resolver la solicitud. "
            f"Modos disponibles: {', '.join(MODOS.keys())}."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "modo": {
                    "type": "string",
                    "description": f"Modo a activar. Opciones: {', '.join(MODOS.keys())}",
                    "enum": list(MODOS.keys())
                }
            },
            "required": ["modo"]
        }
    },
    {
        "name": "crear_agente",
        "description": (
            "Registra un nuevo rol de experto en el sistema tras diseñar su prompt. "
            "Solicita confirmación al usuario antes de crearlo. "
            "Úsalo solo después de construir el prompt con la metodología KIKE."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "nombre_modo": {
                    "type": "string",
                    "description": "Identificador corto sin espacios, ej: 'ml_engineer'"
                },
                "etiqueta": {
                    "type": "string",
                    "description": "Nombre visible del rol, ej: 'Ingeniero de Machine Learning'"
                },
                "contenido_prompt": {
                    "type": "string",
                    "description": "Contenido completo del prompt del nuevo experto"
                }
            },
            "required": ["nombre_modo", "etiqueta", "contenido_prompt"]
        }
    },
    {
        "name": "crear_herramienta",
        "description": (
            "Crea una nueva herramienta Python y la agrega al agente de forma permanente. "
            "El código debe definir una función con el mismo nombre que 'nombre'."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "nombre":      {"type": "string", "description": "Nombre de la herramienta (identificador Python válido)"},
                "descripcion": {"type": "string", "description": "Descripción de qué hace"},
                "parametros":  {"type": "object", "description": "Parámetros: {nombre: {type, description}}"},
                "requeridos":  {"type": "array", "items": {"type": "string"}, "description": "Parámetros obligatorios"},
                "codigo":      {"type": "string", "description": "Código Python autónomo que define la función"}
            },
            "required": ["nombre", "descripcion", "parametros", "requeridos", "codigo"]
        }
    },
    {
        "name": "pedir_confirmacion",
        "description": (
            "Muestra un diálogo de confirmación al usuario en la interfaz gráfica. "
            "Usa esta herramienta SIEMPRE que necesites aprobación del usuario, "
            "en lugar de preguntar en texto en el chat."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "pregunta": {
                    "type": "string",
                    "description": "Descripción de lo que se va a hacer y qué está aprobando el usuario"
                }
            },
            "required": ["pregunta"]
        }
    },
    {
        "name": "solicitar_ruta_proyecto",
        "description": (
            "Configura la carpeta raíz del proyecto. "
            "Si el usuario mencionó una ruta en su mensaje, pásala en 'ruta_sugerida' — "
            "si existe, se configura automáticamente sin mostrar diálogo. "
            "Si no hay ruta sugerida, muestra un selector al usuario. "
            "OBLIGATORIO: llama a esta herramienta ANTES de usar crear_archivo, crear_carpeta o leer_archivo."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "ruta_sugerida": {
                    "type": "string",
                    "description": "Ruta inferida del mensaje del usuario, ej: 'development/cotizacion_lab'. Relativa al directorio del agente o absoluta."
                }
            },
            "required": []
        }
    },
    {
        "name": "crear_carpeta",
        "description": "Crea una carpeta dentro del directorio de proyecto elegido por el usuario.",
        "input_schema": {
            "type": "object",
            "properties": {
                "ruta_relativa": {
                    "type": "string",
                    "description": "Ruta relativa desde la raíz del proyecto, ej: 'src/components'"
                }
            },
            "required": ["ruta_relativa"]
        }
    },
    {
        "name": "crear_archivo",
        "description": "Crea un archivo con contenido completo dentro del directorio de proyecto elegido por el usuario.",
        "input_schema": {
            "type": "object",
            "properties": {
                "ruta_relativa": {
                    "type": "string",
                    "description": "Ruta relativa desde la raíz del proyecto, ej: 'src/App.tsx'"
                },
                "contenido": {
                    "type": "string",
                    "description": "Contenido completo del archivo"
                }
            },
            "required": ["ruta_relativa", "contenido"]
        }
    },
    {
        "name": "listar_archivos",
        "description": "Lista los archivos y carpetas en el directorio del proyecto.",
        "input_schema": {
            "type": "object",
            "properties": {
                "ruta_relativa": {
                    "type": "string",
                    "description": "Sub-ruta a listar (omitir para la raíz del proyecto)"
                }
            },
            "required": []
        }
    },
    {
        "name": "leer_archivo",
        "description": "Lee el contenido completo de un archivo dentro del directorio del proyecto.",
        "input_schema": {
            "type": "object",
            "properties": {
                "ruta_relativa": {
                    "type": "string",
                    "description": "Ruta relativa desde la raíz del proyecto, ej: 'index.html' o 'src/app.py'"
                }
            },
            "required": ["ruta_relativa"]
        }
    },
    {
        "name": "editar_archivo",
        "description": (
            "Reemplaza un fragmento de texto en un archivo existente del proyecto. "
            "Úsala para modificar URLs, variables, bloques de código, etc. sin reescribir el archivo completo."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "ruta_relativa": {
                    "type": "string",
                    "description": "Ruta relativa al archivo dentro del proyecto, ej: 'script.js'"
                },
                "texto_original": {
                    "type": "string",
                    "description": "Texto exacto que se va a reemplazar (debe existir en el archivo)"
                },
                "texto_nuevo": {
                    "type": "string",
                    "description": "Texto que reemplazará al texto_original"
                }
            },
            "required": ["ruta_relativa", "texto_original", "texto_nuevo"]
        }
    },
    {
        "name": "buscar_imagen_web",
        "description": (
            "Busca la URL de una imagen real de un animal o tema en Wikipedia y Wikimedia. "
            "Devuelve URLs de imágenes válidas y gratuitas. "
            "Úsala para obtener URLs reales antes de actualizar imágenes en archivos del proyecto."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "consulta": {
                    "type": "string",
                    "description": "Nombre del animal o tema a buscar, ej: 'Great white shark', 'Bottlenose dolphin'"
                },
                "consulta_es": {
                    "type": "string",
                    "description": "Nombre en español como alternativa, ej: 'Tiburón blanco', 'Delfín mular'"
                }
            },
            "required": ["consulta"]
        }
    },
    {
        "name": "buscar_en_proyecto",
        "description": (
            "Busca un término en TODOS los archivos del proyecto (HTML, CSS, JS, Python, etc.). "
            "Devuelve archivo:línea:contenido para cada coincidencia. "
            "Úsala SIEMPRE antes de diagnosticar un bug de UI — el error puede estar en CSS o HTML, no solo en JS."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "termino": {
                    "type": "string",
                    "description": "Texto a buscar, ej: 'prize-close-btn', 'openTrivia', 'z-index'"
                },
                "extension": {
                    "type": "string",
                    "description": "Filtrar por extensión, ej: '.css', '.js', '.html'. Omitir para buscar en todos."
                }
            },
            "required": ["termino"]
        }
    },
    {
        "name": "buscar_conocimiento",
        "description": (
            "Busca en la base de conocimiento permanente del modo activo. "
            "Úsala cuando necesites recordar información específica del dominio que hayas aprendido antes."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "consulta": {
                    "type": "string",
                    "description": "Término, tema o pregunta a buscar en la base de conocimiento"
                }
            },
            "required": ["consulta"]
        }
    },
    {
        "name": "agregar_conocimiento",
        "description": (
            "Guarda nuevo conocimiento de forma permanente en la base del modo activo. "
            "Úsala cuando aprendas algo valioso: soluciones a problemas, preferencias del usuario, "
            "patrones del dominio, errores comunes y sus correcciones. "
            "El conocimiento persiste entre sesiones y se inyecta automáticamente en tu contexto."
        ),
        "input_schema": {
            "type": "object",
            "properties": {
                "tema": {
                    "type": "string",
                    "description": "Título corto del conocimiento, ej: 'SQL Injection en Django'"
                },
                "contenido": {
                    "type": "string",
                    "description": "Explicación detallada del conocimiento a guardar"
                }
            },
            "required": ["tema", "contenido"]
        }
    },
    {
        "name": "diagnosticar_impresion",
        "description": (
            "Diagnostica automáticamente problemas de impresión y exportación PDF en el proyecto web. "
            "ÚSALA SIEMPRE PRIMERO cuando el usuario reporte: colores incorrectos en PDF, "
            "header blanco al imprimir, fondo que desaparece al guardar PDF, elementos invisibles en PDF. "
            "Escanea JS y CSS del proyecto y retorna un reporte estructurado con todos los problemas encontrados."
        ),
        "input_schema": {
            "type": "object",
            "properties": {},
            "required": []
        }
    }
]

def _tools_activos():
    return [t for t in TOOLS if t["name"] not in TOOLS_DESACTIVADAS]


# ============================================================
# PERSISTENCIA DE MODOS Y HERRAMIENTAS
# ============================================================

def _actualizar_descripcion_cambiar_modo():
    for tool in TOOLS:
        if tool["name"] == "cambiar_modo":
            tool["description"] = f"Cambia el modo del agente. Modos disponibles: {', '.join(MODOS.keys())}."
            tool["input_schema"]["properties"]["modo"]["enum"] = list(MODOS.keys())
            tool["input_schema"]["properties"]["modo"]["description"] = f"Modo a activar. Opciones: {', '.join(MODOS.keys())}"


def cargar_modos_custom():
    if not os.path.exists(CUSTOM_MODES_FILE):
        return
    with open(CUSTOM_MODES_FILE, "r", encoding="utf-8") as f:
        modos = json.load(f)
    for m in modos:
        MODOS[m["nombre_modo"]] = m["archivo"]
        MODOS_ETIQUETAS[m["nombre_modo"]] = m["etiqueta"]
    if modos:
        _actualizar_descripcion_cambiar_modo()
        print(f"[Modos cargados: {', '.join(m['nombre_modo'] for m in modos)}]")


def guardar_modo_custom(nombre_modo, etiqueta, archivo):
    modos = []
    if os.path.exists(CUSTOM_MODES_FILE):
        with open(CUSTOM_MODES_FILE, "r", encoding="utf-8") as f:
            modos = json.load(f)
    entry = {"nombre_modo": nombre_modo, "etiqueta": etiqueta, "archivo": archivo}
    modos = [entry if m["nombre_modo"] == nombre_modo else m for m in modos]
    if nombre_modo not in {m["nombre_modo"] for m in modos}:
        modos.append(entry)
    with open(CUSTOM_MODES_FILE, "w", encoding="utf-8") as f:
        json.dump(modos, f, indent=2, ensure_ascii=False)


def _registrar_tool(nombre, descripcion, input_schema, codigo):
    # _ruta_segura se define más abajo en el módulo; usar lambda para resolución diferida
    # evita NameError cuando cargar_herramientas_custom() se llama antes de esa definición.
    namespace = {
        "RUTA_PROYECTO": RUTA_PROYECTO,
        "_ruta_segura": lambda *a, **kw: _ruta_segura(*a, **kw),
        "_solicitar_confirmacion": _solicitar_confirmacion,
    }
    exec(codigo, namespace)  # noqa: S102
    if nombre not in namespace:
        raise ValueError(f"El código no define una función llamada '{nombre}'")
    HERRAMIENTAS_DINAMICAS[nombre] = namespace[nombre]
    if nombre not in {t["name"] for t in TOOLS}:
        TOOLS.append({"name": nombre, "description": descripcion, "input_schema": input_schema})


def guardar_herramienta_custom(nombre, descripcion, input_schema, codigo):
    tools = []
    if os.path.exists(CUSTOM_TOOLS_FILE):
        with open(CUSTOM_TOOLS_FILE, "r", encoding="utf-8") as f:
            tools = json.load(f)
    entry = {"name": nombre, "description": descripcion, "input_schema": input_schema, "code": codigo}
    tools = [entry if t["name"] == nombre else t for t in tools]
    if nombre not in {t["name"] for t in tools}:
        tools.append(entry)
    with open(CUSTOM_TOOLS_FILE, "w", encoding="utf-8") as f:
        json.dump(tools, f, indent=2, ensure_ascii=False)


def cargar_herramientas_custom():
    if not os.path.exists(CUSTOM_TOOLS_FILE):
        return
    with open(CUSTOM_TOOLS_FILE, "r", encoding="utf-8") as f:
        tools = json.load(f)
    for t in tools:
        try:
            _registrar_tool(t["name"], t["description"], t["input_schema"], t["code"])
            print(f"[Herramienta cargada: {t['name']}]")
        except Exception as e:
            print(f"[Error cargando '{t['name']}': {e}]")


cargar_modos_custom()
cargar_herramientas_custom()


# ============================================================
# IMPLEMENTACIÓN DE HERRAMIENTAS
# ============================================================

def _asegurar_ruta_proyecto():
    """Abre el selector de ruta automáticamente si no hay ruta configurada.
    Retorna None si la ruta ya está configurada o el usuario eligió una.
    Retorna str de error si el usuario canceló o el selector no está disponible."""
    if RUTA_PROYECTO[0]:
        return None
    if _solicitar_ruta[0]:
        ruta = _solicitar_ruta[0]()
        if ruta:
            return None  # _solicitar_ruta[0] ya actualizó RUTA_PROYECTO[0] via callback
        return "El usuario no seleccionó ninguna carpeta de proyecto. No se puede continuar."
    return (
        "Error: no hay ruta de proyecto configurada.\n"
        "Llama solicitar_ruta_proyecto antes de usar herramientas de archivos."
    )


def _ruta_segura(ruta_relativa):
    """Devuelve (ruta_abs, error). error es None si la ruta es válida."""
    err = _asegurar_ruta_proyecto()
    if err:
        return None, err
    base = os.path.abspath(RUTA_PROYECTO[0])
    ruta_rel = (ruta_relativa or "").strip().lstrip("/\\")
    if ruta_rel.startswith(base):
        ruta_rel = ruta_rel[len(base):].lstrip("/\\")
    ruta_abs = os.path.abspath(os.path.join(base, ruta_rel))
    if ruta_abs != base and not ruta_abs.startswith(base + os.sep):
        return None, (
            f"Error: ruta fuera del árbol del proyecto.\n"
            f"  Intentó acceder: {ruta_abs}\n"
            f"  Raíz del proyecto: {base}\n"
            f"  ruta_relativa recibida: '{ruta_relativa}'\n"
            f"  Causa probable: se incluyó la ruta absoluta como relativa, o '..' escapa la raíz.\n"
            f"  Corrección: usa solo nombres de archivo/carpeta relativos (ej: 'src/index.html')."
        )
    return ruta_abs, None


import logging as _logging
import datetime as _datetime

_LOG_PATH = os.path.join(AGENT_DIR, "agent_debug.log")
_log_handler = _logging.FileHandler(_LOG_PATH, encoding="utf-8")
_log_handler.setFormatter(_logging.Formatter("%(asctime)s %(message)s", datefmt="%Y-%m-%d %H:%M:%S"))
_tool_logger = _logging.getLogger("agent.tools")
_tool_logger.setLevel(_logging.DEBUG)
_tool_logger.addHandler(_log_handler)
_tool_logger.propagate = False


def _log_tool(nombre, inputs, result):
    inputs_str = json.dumps(inputs, ensure_ascii=False, default=str)
    if len(inputs_str) > 300:
        inputs_str = inputs_str[:300] + "…"
    result_str = str(result) if result is not None else "<None>"
    if len(result_str) > 400:
        result_str = result_str[:400] + "…"
    _tool_logger.debug("[TOOL] %s | inputs=%s | result=%s", nombre, inputs_str, result_str)


def ejecutar_herramienta(nombre, inputs):
    global SYSTEM_PROMPT, MODO_ACTUAL
    try:
        result = _despachar_herramienta(nombre, inputs)
        _log_tool(nombre, inputs, result)
        return result
    except KeyError as exc:
        msg = (
            f"[ERROR] Herramienta '{nombre}': falta parámetro requerido.\n"
            f"  Campo faltante: {exc}\n"
            f"  Parámetros recibidos: {list(inputs.keys())}\n"
            f"  Causa probable: el modelo omitió un campo obligatorio en el tool call.\n"
            f"  Acción: reintenta incluyendo todos los parámetros requeridos."
        )
        _log_tool(nombre, inputs, msg)
        return msg
    except Exception as exc:
        import traceback as _tb
        msg = (
            f"[ERROR INESPERADO] en herramienta '{nombre}'\n"
            f"  Inputs recibidos: {json.dumps(inputs, ensure_ascii=False, default=str)}\n"
            f"  Tipo: {type(exc).__name__}\n"
            f"  Detalle: {exc}\n"
            f"  Traza:\n{_tb.format_exc()}"
        )
        _log_tool(nombre, inputs, msg)
        return msg


def _despachar_herramienta(nombre, inputs):
    global SYSTEM_PROMPT, MODO_ACTUAL

    if nombre == "calcular":
        operacion = inputs.get("operacion", "")
        try:
            return str(eval(operacion))  # noqa: S307
        except ZeroDivisionError:
            return f"Error en calcular: división por cero en '{operacion}'."
        except SyntaxError as e:
            return f"Error en calcular: sintaxis inválida en '{operacion}' — {e}."
        except Exception as e:
            return f"Error en calcular: no se pudo evaluar '{operacion}' — {type(e).__name__}: {e}."

    if nombre == "cambiar_modo":
        modo = inputs["modo"]
        if modo not in MODOS:
            return f"Modo '{modo}' no existe. Disponibles: {', '.join(MODOS.keys())}"
        if not os.path.exists(MODOS[modo]):
            return f"Archivo de prompt para '{modo}' no encontrado."
        SYSTEM_PROMPT = construir_system_prompt(modo)
        MODO_ACTUAL = modo
        print(f"__MODO__{modo}")
        return f"Modo cambiado a '{MODOS_ETIQUETAS[modo]}'."

    if nombre == "crear_agente":
        nombre_modo    = inputs["nombre_modo"]
        etiqueta       = inputs["etiqueta"]
        contenido      = inputs["contenido_prompt"]

        if nombre_modo in MODOS:
            return f"El modo '{nombre_modo}' ya existe. Elige otro nombre."

        # Pedir permiso al usuario a través de la GUI
        if _solicitar_permiso[0]:
            aprobado = _solicitar_permiso[0](nombre_modo, etiqueta, contenido)
            if not aprobado:
                return f"El usuario rechazó la creación del agente '{etiqueta}'. No se realizó ningún cambio."

        archivo = f"agent_prompt_{nombre_modo}.txt"
        with open(archivo, "w", encoding="utf-8") as f:
            f.write(contenido)

        MODOS[nombre_modo] = archivo
        MODOS_ETIQUETAS[nombre_modo] = etiqueta
        guardar_modo_custom(nombre_modo, etiqueta, archivo)
        _actualizar_descripcion_cambiar_modo()

        return (
            f"Agente '{etiqueta}' creado y registrado exitosamente.\n"
            f"Archivo: {archivo}\n"
            f"Usa cambiar_modo('{nombre_modo}') para activarlo."
        )

    if nombre == "crear_herramienta":
        try:
            input_schema = {
                "type": "object",
                "properties": inputs["parametros"],
                "required": inputs.get("requeridos", list(inputs["parametros"].keys()))
            }
            _registrar_tool(inputs["nombre"], inputs["descripcion"], input_schema, inputs["codigo"])
            guardar_herramienta_custom(inputs["nombre"], inputs["descripcion"], input_schema, inputs["codigo"])
            return (
                f"Herramienta '{inputs['nombre']}' creada y registrada exitosamente.\n"
                f"ACCIÓN REQUERIDA: continúa AHORA con la tarea original del usuario "
                f"usando esta nueva herramienta. No pidas confirmación ni hagas preguntas — "
                f"llama la herramienta directamente."
            )
        except Exception as e:
            return f"Error al crear herramienta: {e}"

    if nombre == "pedir_confirmacion":
        pregunta = inputs.get("pregunta", "¿Confirmas esta acción?")
        if _solicitar_confirmacion[0]:
            aprobado = _solicitar_confirmacion[0](pregunta)
            return "El usuario aprobó. Procede." if aprobado else "El usuario rechazó. No realices la acción."
        return "Confirmación no disponible. Procede."

    if nombre == "solicitar_ruta_proyecto":
        ruta_sugerida = inputs.get("ruta_sugerida", "").strip()
        if ruta_sugerida:
            if os.path.isabs(ruta_sugerida):
                ruta_resuelta = ruta_sugerida
            else:
                ruta_resuelta = os.path.join(AGENT_DIR, ruta_sugerida)
            ruta_resuelta = os.path.abspath(ruta_resuelta)
            if os.path.isdir(ruta_resuelta):
                RUTA_PROYECTO[0] = ruta_resuelta
                return (
                    f"Ruta del proyecto configurada automáticamente: {ruta_resuelta}\n"
                    f"IMPORTANTE: ruta_relativa es RELATIVA a esa carpeta raíz.\n"
                    f"  - Para listar la raíz: ruta_relativa=\"\"\n"
                    f"  - Para un archivo en la raíz: ruta_relativa=\"index.html\"\n"
                    f"  - Para una subcarpeta: ruta_relativa=\"src/App.tsx\"\n"
                    f"  NUNCA incluyas '{ruta_resuelta}' en ruta_relativa.\n"
                    f"Procede AHORA con la tarea. NO hagas más preguntas."
                )
        if _solicitar_ruta[0]:
            ruta = _solicitar_ruta[0]()
            if ruta:
                os.makedirs(ruta, exist_ok=True)
                if BACKEND == "claude-code":
                    instruccion_crear = (
                        f"Procede AHORA a crear los archivos usando bloques <file>:\n"
                        f'<file path="nombre.ext">contenido completo</file>'
                    )
                else:
                    instruccion_crear = (
                        f"Procede AHORA a crear cada archivo con tool calls:\n"
                        f'<tool_call>{{"name": "crear_archivo", "inputs": {{"ruta_relativa": "nombre.ext", "contenido": "contenido completo"}}}}</tool_call>'
                    )
                return (
                    f"Ruta del proyecto APROBADA por el usuario: {ruta}\n"
                    f"IMPORTANTE: ruta_relativa es RELATIVA a esa carpeta raíz.\n"
                    f"  - Para listar la raíz: ruta_relativa=\"\"\n"
                    f"  - Para un archivo en la raíz: ruta_relativa=\"index.html\"\n"
                    f"  - Para una subcarpeta: ruta_relativa=\"src/App.tsx\"\n"
                    f"  NUNCA incluyas '{ruta}' ni ninguna parte de la ruta absoluta en ruta_relativa.\n"
                    f"El usuario ya dio su aprobación al seleccionar esta carpeta.\n"
                    f"{instruccion_crear}\n"
                    f"NO hagas más preguntas. NO pidas más permisos."
                )
            return "El usuario no seleccionó ninguna ruta. No puedes crear archivos sin una ruta destino."
        return "Error: selector de ruta no disponible en este contexto."

    if nombre == "crear_carpeta":
        ruta_abs, err = _ruta_segura(inputs.get("ruta_relativa", ""))
        if err:
            return err
        try:
            os.makedirs(ruta_abs, exist_ok=True)
        except OSError as e:
            return (
                f"Error al crear carpeta '{ruta_abs}'.\n"
                f"  Causa: {type(e).__name__}: {e}\n"
                f"  Verifica permisos de escritura en el directorio padre."
            )
        return f"Carpeta creada: {ruta_abs}"

    if nombre == "crear_archivo":
        ruta_abs, err = _ruta_segura(inputs["ruta_relativa"])
        if err:
            return err
        if _solicitar_confirmacion[0]:
            contenido_preview = inputs.get("contenido", "")[:300].replace("\n", "↵")
            n_chars = len(inputs.get("contenido", ""))
            pregunta = (
                f"El agente quiere CREAR el archivo '{inputs['ruta_relativa']}' "
                f"({n_chars} caracteres).\n\n"
                f"Vista previa:\n{contenido_preview}"
                f"{'…' if n_chars > 300 else ''}"
            )
            aprobado = _solicitar_confirmacion[0](pregunta)
            if not aprobado:
                return f"El usuario rechazó crear '{inputs['ruta_relativa']}'. No se realizó ningún cambio."
        try:
            os.makedirs(os.path.dirname(ruta_abs) or RUTA_PROYECTO[0], exist_ok=True)
            with open(ruta_abs, "w", encoding="utf-8") as f:
                f.write(inputs["contenido"])
        except OSError as e:
            return (
                f"Error al escribir '{ruta_abs}'.\n"
                f"  Causa: {type(e).__name__}: {e}\n"
                f"  Intentó escribir: {len(inputs.get('contenido',''))} caracteres.\n"
                f"  Verifica permisos y que el directorio padre exista."
            )
        return f"Archivo creado: {ruta_abs}"

    if nombre == "leer_archivo":
        ruta_abs, err = _ruta_segura(inputs.get("ruta_relativa", ""))
        if err:
            return err
        if not os.path.isfile(ruta_abs):
            return (
                f"Error: '{inputs.get('ruta_relativa')}' no existe o no es un archivo.\n"
                f"Usa listar_archivos con ruta_relativa='' para ver los archivos disponibles."
            )
        # --- Auto-conversión de .docx a texto plano ---
        if ruta_abs.lower().endswith(".docx"):
            try:
                import docx as _docx_mod
                doc = _docx_mod.Document(ruta_abs)
                partes = []
                for p in doc.paragraphs:
                    if p.text.strip():
                        partes.append(p.text)
                # Incluir tablas
                for table in doc.tables:
                    for row in table.rows:
                        fila = " | ".join(c.text.strip() for c in row.cells if c.text.strip())
                        if fila:
                            partes.append(fila)
                contenido = "\n".join(partes)
                if not contenido.strip():
                    return (
                        f"[DOCX VACÍO] '{inputs.get('ruta_relativa')}' es un .docx sin texto extraíble."
                    )
                # Guardar .txt al lado del .docx
                ruta_txt = ruta_abs[:-5] + ".txt"
                with open(ruta_txt, "w", encoding="utf-8") as f:
                    f.write(contenido)
                ruta_rel = inputs.get("ruta_relativa", "")
                ruta_txt_rel = ruta_rel[:-5] + ".txt" if ruta_rel.lower().endswith(".docx") else ruta_txt
                return (
                    f"[DOCX → TXT] Convertido automáticamente.\n"
                    f"Texto guardado en: {ruta_txt_rel}\n"
                    f"{'─'*50}\n"
                    f"{contenido}"
                )
            except ImportError:
                return (
                    "Error: 'python-docx' no instalado.\n"
                    "Ejecuta: pip3 install python-docx"
                )
            except Exception as e:
                return f"Error al leer .docx '{ruta_abs}': {type(e).__name__}: {e}"
        # --- Lectura normal de texto plano ---
        try:
            with open(ruta_abs, "r", encoding="utf-8") as f:
                contenido = f.read()
            if not contenido.strip():
                return (
                    f"[ARCHIVO VACÍO] '{inputs.get('ruta_relativa')}' existe pero tiene 0 bytes de contenido.\n"
                    f"  Causa probable: el archivo fue creado pero nunca se escribió contenido en él.\n"
                    f"  Acción: verifica cómo fue creado el archivo o pide al usuario que lo rellene."
                )
            return contenido
        except UnicodeDecodeError:
            return (
                f"Error: '{inputs.get('ruta_relativa')}' no es un archivo de texto legible (binario o encoding no UTF-8).\n"
                f"  Si es un .docx, asegúrate de que la extensión sea .docx."
            )
        except OSError as e:
            return f"Error al leer '{ruta_abs}': {type(e).__name__}: {e}"

    if nombre == "editar_archivo":
        ruta_abs, err = _ruta_segura(inputs["ruta_relativa"])
        if err:
            return err
        if not os.path.isfile(ruta_abs):
            return (
                f"Error: '{inputs['ruta_relativa']}' no existe o no es un archivo.\n"
                f"Usa listar_archivos con ruta_relativa='' para ver los archivos disponibles."
            )
        if _solicitar_confirmacion[0]:
            orig_preview = inputs["texto_original"][:200].replace("\n", "↵")
            nuevo_preview = inputs["texto_nuevo"][:200].replace("\n", "↵")
            pregunta = (
                f"El agente quiere EDITAR '{inputs['ruta_relativa']}'.\n\n"
                f"Reemplazar:\n{orig_preview}"
                f"{'…' if len(inputs['texto_original']) > 200 else ''}\n\n"
                f"Con:\n{nuevo_preview}"
                f"{'…' if len(inputs['texto_nuevo']) > 200 else ''}"
            )
            aprobado = _solicitar_confirmacion[0](pregunta)
            if not aprobado:
                return f"El usuario rechazó editar '{inputs['ruta_relativa']}'. No se realizó ningún cambio."
        try:
            with open(ruta_abs, "r", encoding="utf-8") as f:
                contenido = f.read()
        except OSError as e:
            return f"Error al leer '{ruta_abs}': {type(e).__name__}: {e}"
        texto_original = inputs["texto_original"]
        if texto_original not in contenido:
            # Intentar match parcial con los primeros 60 chars para dar contexto útil
            partial = texto_original[:60]
            idx = contenido.find(partial)
            if idx != -1:
                start = max(0, idx - 100)
                end = min(len(contenido), idx + 300)
                contexto = contenido[start:end].replace("\n", "↵")
                return (
                    f"Error: texto_original no encontrado EXACTAMENTE en '{inputs['ruta_relativa']}'. "
                    f"Los primeros 60 chars SÍ se encontraron en la línea ~{contenido[:idx].count(chr(10))+1}, "
                    f"pero el texto completo no coincide (diferencia en espacios, tabs o saltos de línea).\n"
                    f"  Buscado: {repr(texto_original[:120])}\n"
                    f"  Contexto real alrededor del match parcial:\n  {contexto}\n"
                    f"  → Llama leer_archivo('{inputs['ruta_relativa']}') para copiar el texto exacto."
                )
            else:
                preview = contenido[:400].replace("\n", "↵")
                return (
                    f"Error: texto_original no encontrado en '{inputs['ruta_relativa']}'.\n"
                    f"  Buscado: {repr(texto_original[:120])}\n"
                    f"  El texto no existe en el archivo. "
                    f"Es posible que ya fue modificado o el texto fue generado de memoria.\n"
                    f"  Inicio del archivo: {preview}\n"
                    f"  → Llama leer_archivo('{inputs['ruta_relativa']}') para ver el contenido exacto y copiar el texto a editar."
                )
        nuevo_contenido = contenido.replace(texto_original, inputs["texto_nuevo"], 1)
        try:
            with open(ruta_abs, "w", encoding="utf-8") as f:
                f.write(nuevo_contenido)
        except OSError as e:
            return f"Error al escribir '{ruta_abs}': {type(e).__name__}: {e}"
        return f"Archivo editado: {inputs['ruta_relativa']} — reemplazo aplicado."

    if nombre == "buscar_imagen_web":
        import urllib.request
        import urllib.parse
        import ssl
        try:
            import certifi
            _ssl_ctx = ssl.create_default_context(cafile=certifi.where())
        except ImportError:
            _ssl_ctx = ssl.create_default_context()
        consulta = inputs.get("consulta", "")
        consulta_es = inputs.get("consulta_es", "")
        resultados = []

        import re as _re
        import urllib.error as _uerr

        def _fetch_wikipedia(lang, titulo):
            titulo_enc = urllib.parse.quote(titulo.replace(" ", "_"))
            url = f"https://{lang}.wikipedia.org/api/rest_v1/page/summary/{titulo_enc}"
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            try:
                with urllib.request.urlopen(req, timeout=8, context=_ssl_ctx) as r:
                    data = json.loads(r.read().decode())
                img = data.get("thumbnail", {}).get("source", "")
                if img:
                    img = _re.sub(r'/\d+px-', '/800px-', img)
                return img, data.get("title", titulo), None
            except _uerr.HTTPError as he:
                return "", titulo, f"HTTP {he.code}"
            except Exception as ex:
                return "", titulo, str(ex)

        errores = []
        intentos = [(consulta, "en"), (consulta_es or consulta, "es")]
        for term, lang in intentos:
            if not term:
                continue
            img_url, titulo, err = _fetch_wikipedia(lang, term)
            if img_url:
                resultados.append({"url": img_url, "fuente": f"Wikipedia ({lang}) — {titulo}"})
            elif err:
                errores.append(f"Wikipedia ({lang}/{term}): {err}")

        if not resultados:
            detalle = ("\n  Errores: " + "; ".join(errores)) if errores else ""
            return (
                f"No se encontraron imágenes para '{consulta}'.{detalle}\n"
                f"  Intenta con el nombre exacto del artículo en Wikipedia en inglés.\n"
                f"  Ejemplos válidos: 'Great white shark', 'Bottlenose dolphin', 'Humpback whale', 'Sea star'.\n"
                f"  O usa editar_archivo con una URL de imagen directamente."
            )
        lines = [f"Imágenes encontradas para '{consulta}':"]
        for r in resultados:
            lines.append(f"  URL: {r['url']}")
            lines.append(f"  Fuente: {r['fuente']}")
        return "\n".join(lines)

    if nombre == "listar_archivos":
        err = _asegurar_ruta_proyecto()
        if err:
            return err
        ruta_rel = inputs.get("ruta_relativa", "").strip().lstrip("/\\")
        # Si el modelo pasó la ruta absoluta del proyecto como relativa, ignorarla
        base = os.path.abspath(RUTA_PROYECTO[0])
        if ruta_rel and os.path.abspath(ruta_rel) == base:
            ruta_rel = ""
        if ruta_rel and ruta_rel.startswith(base):
            ruta_rel = ruta_rel[len(base):].lstrip("/\\")
        ruta_abs = os.path.join(base, ruta_rel) if ruta_rel else base
        ruta_abs = os.path.abspath(ruta_abs)
        if not os.path.exists(ruta_abs):
            return (
                f"Error en listar_archivos: la ruta no existe.\n"
                f"  Ruta intentada: {ruta_abs}\n"
                f"  ruta_relativa recibida: '{inputs.get('ruta_relativa', '')}'\n"
                f"  Raíz del proyecto: {base}\n"
                f"  Causa probable: subcarpeta no creada aún, o ruta_relativa incorrecta.\n"
                f"  Para listar la raíz usa ruta_relativa='' (cadena vacía)."
            )
        items = sorted(os.listdir(ruta_abs))
        if not items:
            return "(carpeta vacía)"
        return "\n".join(
            f"{'[dir]  ' if os.path.isdir(os.path.join(ruta_abs, i)) else '[arch] '}{i}"
            for i in items
        )

    if nombre == "buscar_en_proyecto":
        err = _asegurar_ruta_proyecto()
        if err:
            return err
        termino = inputs.get("termino", "").strip()
        if not termino:
            return "Se requiere 'termino'."
        extension = inputs.get("extension", "").strip().lower()
        base = os.path.abspath(RUTA_PROYECTO[0])
        resultados = []
        EXTENSIONES_TEXTO = {".html", ".css", ".js", ".ts", ".jsx", ".tsx", ".py", ".json", ".md", ".txt", ".env", ".yaml", ".yml"}
        for root, dirs, files in os.walk(base):
            dirs[:] = [d for d in dirs if not d.startswith(".") and d != "node_modules"]
            for fname in sorted(files):
                if extension and not fname.endswith(extension):
                    continue
                ext = os.path.splitext(fname)[1].lower()
                if not extension and ext not in EXTENSIONES_TEXTO:
                    continue
                fpath = os.path.join(root, fname)
                rel = os.path.relpath(fpath, base)
                try:
                    with open(fpath, "r", encoding="utf-8", errors="ignore") as f:
                        for i, line in enumerate(f, 1):
                            if termino.lower() in line.lower():
                                resultados.append(f"{rel}:{i}: {line.rstrip()}")
                except OSError:
                    pass
        if not resultados:
            return f"'{termino}' no encontrado en ningún archivo del proyecto."
        if len(resultados) > 80:
            resultados = resultados[:80]
            resultados.append(f"[... resultado truncado a 80 líneas]")
        return "\n".join(resultados)

    if nombre == "buscar_conocimiento":
        consulta = inputs.get("consulta", "").lower()
        entries = cargar_conocimiento(MODO_ACTUAL)
        if not entries:
            return "La base de conocimiento está vacía para este modo."
        resultados = [
            e for e in entries
            if consulta in e["tema"].lower() or consulta in e["contenido"].lower()
        ]
        if not resultados:
            return f"No se encontró nada relacionado con '{inputs.get('consulta', '')}'."
        return "\n\n".join(f"[{e['tema']}]\n{e['contenido']}" for e in resultados)

    if nombre == "agregar_conocimiento":
        tema     = inputs.get("tema", "").strip()
        contenido = inputs.get("contenido", "").strip()
        if not tema or not contenido:
            return "Se requieren 'tema' y 'contenido'."
        entries = cargar_conocimiento(MODO_ACTUAL)
        existente = next((e for e in entries if e["tema"].lower() == tema.lower()), None)
        if existente:
            existente["contenido"] = contenido
        else:
            entries.append({"tema": tema, "contenido": contenido})
        guardar_conocimiento(MODO_ACTUAL, entries)
        SYSTEM_PROMPT = construir_system_prompt(MODO_ACTUAL)
        return f"Conocimiento '{tema}' guardado para el modo '{MODO_ACTUAL}'."

    if nombre == "diagnosticar_impresion":
        err = _asegurar_ruta_proyecto()
        if err:
            return err
        base = os.path.abspath(RUTA_PROYECTO[0])
        reporte = ["=== DIAGNÓSTICO DE IMPRESIÓN/PDF ===\n"]

        # --- 1. Escanear archivos JS ---
        reporte.append("## JavaScript — interferencia con impresión")
        js_encontrado = False
        for root_dir, dirs, files in os.walk(base):
            dirs[:] = [d for d in dirs if not d.startswith(".") and d != "node_modules"]
            for fname in sorted(files):
                if not fname.endswith(".js"):
                    continue
                fpath = os.path.join(root_dir, fname)
                rel = os.path.relpath(fpath, base)
                try:
                    with open(fpath, "r", encoding="utf-8", errors="ignore") as f:
                        lines = f.readlines()
                    tiene_guard = any(
                        "matchMedia('print')" in l or 'matchMedia("print")' in l
                        for l in lines
                    )
                    tiene_beforeprint = any("beforeprint" in l for l in lines)
                    for i, line in enumerate(lines, 1):
                        ls = line.strip()
                        if "IntersectionObserver" in line:
                            reporte.append(f"  ⚠️  [{rel}:{i}] IntersectionObserver detectado")
                            if not tiene_guard:
                                reporte.append(f"       → Falta guardia: if (!window.matchMedia('print').matches)")
                            if not tiene_beforeprint:
                                reporte.append(f"       → Falta listener beforeprint para resetear opacity/transform")
                            js_encontrado = True
                        if "style.opacity" in line and "'0'" in line:
                            reporte.append(f"  ⚠️  [{rel}:{i}] opacity='0' sin guardia de impresión: {ls}")
                            js_encontrado = True
                        if "body.style.opacity" in line:
                            reporte.append(f"  ⚠️  [{rel}:{i}] body.style.opacity: {ls}")
                            if not tiene_beforeprint:
                                reporte.append(f"       → Falta: window.addEventListener('beforeprint', () => {{ document.body.style.opacity='1'; }})")
                            js_encontrado = True
                        if "beforeprint" in line:
                            reporte.append(f"  ✅  [{rel}:{i}] beforeprint listener presente")
                        if "matchMedia('print')" in line or 'matchMedia("print")' in line:
                            reporte.append(f"  ✅  [{rel}:{i}] Guardia de impresión presente")
                except OSError:
                    pass
        if not js_encontrado:
            reporte.append("  Sin patrones de interferencia JS detectados.")

        # --- 2. Escanear archivos CSS ---
        reporte.append("\n## CSS — reglas de impresión")
        for root_dir, dirs, files in os.walk(base):
            dirs[:] = [d for d in dirs if not d.startswith(".") and d != "node_modules"]
            for fname in sorted(files):
                if not fname.endswith(".css"):
                    continue
                fpath = os.path.join(root_dir, fname)
                rel = os.path.relpath(fpath, base)
                try:
                    with open(fpath, "r", encoding="utf-8", errors="ignore") as f:
                        lines = f.readlines()

                    media_print_lines = []
                    for i, line in enumerate(lines, 1):
                        if "@media print" in line.lower():
                            media_print_lines.append(i)

                    if len(media_print_lines) > 1:
                        reporte.append(f"  ⚠️  [{rel}] {len(media_print_lines)} bloques @media print"
                                       f" (líneas {media_print_lines}) → consolidar en uno")
                    elif len(media_print_lines) == 1:
                        reporte.append(f"  ✅  [{rel}] 1 bloque @media print (línea {media_print_lines[0]})")
                    else:
                        reporte.append(f"  ℹ️  [{rel}] Sin @media print")

                    for i, line in enumerate(lines, 1):
                        ls = line.strip()
                        if "preserve" in line and "print-color-adjust" in line.lower():
                            reporte.append(f"  ⚠️  [{rel}:{i}] PRESERVE: {ls}")
                            reporte.append(f"       → Cambiar a: print-color-adjust: exact !important")
                        if ("background: none" in line.lower() or "background:none" in line.lower()):
                            ctx_start = max(0, i - 12)
                            ctx = "".join(lines[ctx_start:i - 1])
                            # Check if the selector group before this line includes a root element
                            # (selector without child combinator as root)
                            selector_lines = []
                            for j in range(i - 2, ctx_start - 1, -1):
                                sl = lines[j].strip() if j < len(lines) else ""
                                if sl.endswith("{"):
                                    break
                                if sl.endswith(",") or (not sl.startswith(".") is False):
                                    selector_lines.append(sl)
                            has_root_selector = any(
                                l.strip().endswith(",") and " " not in l.strip().rstrip(",")
                                for l in lines[ctx_start:i - 1]
                                if l.strip().endswith(",")
                            )
                            reporte.append(f"  ⚠️  [{rel}:{i}] background:none — verificar si el selector incluye el contenedor raíz")
                            reporte.append(f"       → Si '.elemento,' aparece en el grupo, ese elemento perderá su fondo")
                            reporte.append(f"       → Contexto: líneas {ctx_start + 1}–{i}")
                except OSError:
                    pass

        reporte.append("\n## Causa más probable")
        reporte.append("1. print-color-adjust: preserve → cambiar a exact")
        reporte.append("2. IntersectionObserver sin guardia → envolver con matchMedia check + beforeprint")
        reporte.append("3. background:none en selector que incluye contenedor → separar selectores")
        reporte.append("4. Múltiples @media print → consolidar en uno")
        reporte.append("\nRevisa cada ⚠️ arriba y corrige en ese orden.")
        return "\n".join(reporte)

    if nombre in HERRAMIENTAS_DINAMICAS:
        try:
            # Resolver parámetros "ruta*" relativos contra RUTA_PROYECTO[0]
            if RUTA_PROYECTO[0]:
                base = os.path.abspath(RUTA_PROYECTO[0])
                inputs_resueltos = {}
                for k, v in inputs.items():
                    if k.startswith("ruta") and isinstance(v, str) and v and not os.path.isabs(v):
                        ruta_resuelta = os.path.abspath(os.path.join(base, v))
                        # Bloquear rutas que escapan del árbol del proyecto
                        if ruta_resuelta.startswith(base):
                            inputs_resueltos[k] = ruta_resuelta
                        else:
                            return (
                                f"Error: ruta '{v}' escapa del árbol del proyecto.\n"
                                f"  Raíz: {base}\n"
                                f"  Resuelto: {ruta_resuelta}"
                            )
                    else:
                        inputs_resueltos[k] = v
                inputs = inputs_resueltos

            tool_def = next((t for t in TOOLS if t["name"] == nombre), {})
            desc = tool_def.get("description", "").lower()
            palabras_escritura = ["crear", "guardar", "escribir", "create", "write", "save", "editar", "edit", "generar", "generate"]
            es_escritura = any(p in desc for p in palabras_escritura)
            if es_escritura and _solicitar_confirmacion[0]:
                args_preview = json.dumps(inputs, ensure_ascii=False, default=str)
                if len(args_preview) > 400:
                    args_preview = args_preview[:400] + "…"
                aprobado = _solicitar_confirmacion[0](
                    f"La herramienta '{nombre}' quiere ejecutarse con:\n{args_preview}\n\n¿Aprobar?"
                )
                if not aprobado:
                    return f"El usuario rechazó la ejecución de '{nombre}'. No se realizó ningún cambio."
            return str(HERRAMIENTAS_DINAMICAS[nombre](**inputs))
        except TypeError as e:
            return (
                f"Error en herramienta dinámica '{nombre}': firma incompatible.\n"
                f"  Inputs recibidos: {list(inputs.keys())}\n"
                f"  Detalle: {e}"
            )
        except Exception as e:
            return (
                f"Error ejecutando herramienta dinámica '{nombre}'.\n"
                f"  Inputs: {json.dumps(inputs, ensure_ascii=False, default=str)}\n"
                f"  {type(e).__name__}: {e}"
            )

    return (
        f"Herramienta '{nombre}' no encontrada.\n"
        f"  Herramientas disponibles: {[t['name'] for t in _tools_activos()]}\n"
        f"  Causa probable: nombre mal escrito o herramienta desactivada en la UI."
    )


# ============================================================
# LOOPS POR BACKEND
# ============================================================

def _formatear_herramientas_texto():
    """Convierte TOOLS a descripción de texto para el prompt."""
    lineas = [
        "## CÓMO ENTREGAR ARCHIVOS — SIGUE ESTE EJEMPLO EXACTAMENTE",
        "",
        "Cuando el usuario pide crear o modificar archivos, incluye el código completo así:",
        "",
        "**`index.html`**",
        "```html",
        "<!DOCTYPE html>",
        "<html><body><h1>Ejemplo</h1></body></html>",
        "```",
        "",
        "**`style.css`**",
        "```css",
        "body { font-family: sans-serif; }",
        "```",
        "",
        "Eso es todo. No expliques permisos, no menciones terminal ni diálogos.",
        "No hay Opción A / Opción B. Solo pon el código de cada archivo en tu respuesta.",
        "",
        "## HERRAMIENTAS DISPONIBLES — USO OBLIGATORIO",
        "",
        "Para ejecutar cualquier acción usa este formato. NUNCA describas que la harás — hazla:",
        "",
        "Ejemplo cambiar experto:",
        '<tool_call>{"name": "cambiar_modo", "inputs": {"modo": "fullstack"}}</tool_call>',
        "",
        "Ejemplo pedir ruta:",
        '<tool_call>{"name": "solicitar_ruta_proyecto", "inputs": {}}</tool_call>',
        "",
        "Ejemplo crear archivo:",
        '<tool_call>{"name": "crear_archivo", "inputs": {"ruta_relativa": "index.html", "contenido": "<!DOCTYPE html>..."}}</tool_call>',
        "",
        "REGLAS ESTRICTAS:",
        "- Todo el JSON en una sola línea dentro de las etiquetas <tool_call>.",
        "- NUNCA escribas 'llamaré a X' ni 'voy a ejecutar Y' — ejecútalo directamente.",
        "- NUNCA simules el resultado — espera la respuesta del sistema.",
        "- Sin ruta de proyecto: llama solicitar_ruta_proyecto ANTES de crear archivos.",
        "- Tarea técnica: llama cambiar_modo con el experto correcto PRIMERO.",
        "",
    ]
    for tool in _tools_activos():
        lineas.append(f"**{tool['name']}**: {tool['description']}")
        props = tool["input_schema"].get("properties", {})
        reqs  = tool["input_schema"].get("required", [])
        for pname, pdef in props.items():
            marca = "requerido" if pname in reqs else "opcional"
            lineas.append(f"  · {pname} ({marca}): {pdef.get('description', '')}")
        lineas.append("")
    return "\n".join(lineas)


def _extraer_tool_calls(texto):
    calls = []
    # Acepta con o sin tag de cierre (algunos modelos lo omiten)
    patron = r"<tool_call>(.*?)(?:</tool_call>|$)"
    for m in re.finditer(patron, texto, re.DOTALL):
        raw = m.group(1).strip()
        if not raw:
            continue
        try:
            calls.append(json.loads(raw))
        except Exception:
            pass
    return calls


def _extraer_archivos(texto):
    """Extrae bloques <file path="...">...</file> del texto."""
    archivos = []
    for m in re.finditer(r'<file\s+path="([^"]+)">(.*?)</file>', texto, re.DOTALL):
        archivos.append({
            "path": m.group(1).strip(),
            "content": m.group(2).lstrip("\n")
        })
    return archivos


def _extraer_archivos_markdown(texto):
    """Fallback: extrae archivos de respuestas markdown **`file.ext`** + bloque código."""
    archivos = []
    patron = re.compile(
        r'\*\*`([a-zA-Z0-9_.\-/]+\.[a-zA-Z]{1,10})`\*\*\s*\n+```[a-z]*\n(.*?)```',
        re.DOTALL
    )
    for m in patron.finditer(texto):
        contenido = m.group(2)
        if contenido.strip():
            archivos.append({"path": m.group(1).strip(), "content": contenido})
    return archivos


_FILTRO_PERMISO_RE = re.compile(
    r'[^\n]*(necesitas aprobar|haz clic en\s+\*{0,2}Allow|cuando aparezca.{0,40}aviso'
    r'|permiso de escritura en Claude|diálogo del sistema|click.{0,10}Allow'
    r'|approve.{0,20}permission|Para poder crear los archivos necesito que apruebes'
    r'|Ve a la terminal|pre-autorizar escritura|bash setup\.sh|script de instalación'
    r'|diálogos de permiso de Claude|aparecen en la terminal)[^\n]*\n?',
    re.IGNORECASE
)

_META_EXPLICACION_RE = re.compile(
    r'terminal donde está corriendo|settings\.json|setup\.sh'
    r'|Opción [ABC] —|¿Cuál prefieres|pre-autorizar|diálogos de permiso de Claude'
    r'|bash setup|script de instalación|escribe `?y`? y presiona'
    r'|no tengo permisos|permisos de escritura|no puedo (crear|escribir|guardar)'
    r'|para que lo guardes|guárdalo (tú|manualmente)|te entrego el código',
    re.IGNORECASE
)

def correr_agente_anthropic(mensaje_usuario):
    global _historial
    _historial.append({"role": "user", "content": mensaje_usuario})
    messages = [{"role": e["role"], "content": e["content"]} for e in _historial]
    while True:
        if _detener_agente.is_set():
            break
        response = client.messages.create(
            model=MODELO, max_tokens=4096, thinking={"type": "adaptive"},
            system=SYSTEM_PROMPT, tools=_tools_activos(), messages=messages
        )
        if response.stop_reason == "end_turn":
            for block in response.content:
                if hasattr(block, "text"):
                    print("\nAgente:", block.text)
                    _historial.append({"role": "assistant", "content": block.text})
            break
        if response.stop_reason == "tool_use":
            messages.append({"role": "assistant", "content": response.content})
            tool_results = []
            for block in response.content:
                if block.type == "tool_use":
                    print(f"[Herramienta: {block.name}]")
                    resultado = ejecutar_herramienta(block.name, block.input)
                    tool_results.append({"type": "tool_result", "tool_use_id": block.id, "content": resultado})
            messages.append({"role": "user", "content": tool_results})


def correr_agente_ollama(mensaje_usuario):
    global _historial
    tools_text = _formatear_herramientas_texto()
    system_con_tools = f"{SYSTEM_PROMPT}\n\n{tools_text}"
    _historial.append({"role": "user", "content": mensaje_usuario})
    messages = [{"role": "system", "content": system_con_tools}]
    messages += [{"role": e["role"], "content": e["content"]} for e in _historial]

    MAX_ITER = 10
    respuesta_final = ""
    for _ in range(MAX_ITER):
        if _detener_agente.is_set():
            break
        response = client.chat(model=MODELO, messages=messages)
        content = response.message.content or ""
        messages.append({"role": "assistant", "content": content})

        tool_calls = _extraer_tool_calls(content)
        if not tool_calls:
            texto_visible = re.sub(r"<tool_call>.*?(?:</tool_call>|$)", "", content, flags=re.DOTALL).strip()
            respuesta_final = texto_visible or content
            print("\nAgente:", respuesta_final)
            break

        resultados = []
        for tc in tool_calls:
            nombre = tc.get("name", "")
            inputs_tc = tc.get("inputs", {})
            print(f"[Herramienta: {nombre}]")
            resultado = ejecutar_herramienta(nombre, inputs_tc)
            if nombre == "leer_archivo":
                print(f"[→ {inputs_tc.get('ruta_relativa', '?')}]")
            else:
                print(f"[→ {resultado}]")
            resultados.append(f"{nombre}: {resultado}")

        messages.append({
            "role": "user",
            "content": "[RESULTADO DE HERRAMIENTAS]:\n" + "\n".join(resultados)
        })

    if respuesta_final:
        _historial.append({"role": "assistant", "content": respuesta_final})




_historial: list = []
_detener_agente = threading.Event()


def compactar_historial():
    global _historial

    if not _historial:
        return "Historial vacío."

    fragmentos = []
    for e in _historial[-30:]:
        rol = e["role"].upper()
        contenido = str(e.get("content") or "")[:400]
        fragmentos.append(f"[{rol}]: {contenido}")
    historial_texto = "\n".join(fragmentos)

    prompt_resumen = (
        "Resume esta conversación en máximo 250 palabras. "
        "Incluye: (1) objetivo del usuario, (2) qué ya se creó o hizo, "
        "(3) estado actual de la tarea, (4) decisiones o preferencias clave. "
        "Solo el resumen, sin explicaciones adicionales.\n\n"
        f"HISTORIAL:\n{historial_texto}"
    )

    resumen = None
    try:
        if BACKEND == "zhipu" and client:
            resp = client.chat.completions.create(
                model=MODELO,
                messages=[{"role": "user", "content": prompt_resumen}],
            )
            resumen = resp.choices[0].message.content or ""
        elif BACKEND == "anthropic" and client:
            resp = client.messages.create(
                model=MODELO, max_tokens=512,
                messages=[{"role": "user", "content": prompt_resumen}]
            )
            resumen = resp.content[0].text if resp.content else ""
        elif BACKEND == "ollama" and client:
            resp = client.chat(
                model=MODELO,
                messages=[{"role": "user", "content": prompt_resumen}]
            )
            resumen = resp.message.content or ""
    except Exception as e:
        resumen = f"[Resumen automático falló: {e}]"

    if not resumen:
        resumen = "[Sin resumen generado]"

    ruta_info = f"\nRuta del proyecto activa: {RUTA_PROYECTO[0]}" if RUTA_PROYECTO[0] else ""
    modo_info = f"\nModo activo: {MODOS_ETIQUETAS.get(MODO_ACTUAL, MODO_ACTUAL)}"

    contexto_compactado = (
        f"[CONTEXTO PREVIO — conversación compactada]{modo_info}{ruta_info}\n\n"
        f"{resumen}"
    )

    _historial.clear()
    _historial.append({"role": "user", "content": contexto_compactado})
    _historial.append({"role": "assistant", "content": "Contexto cargado. Continúo desde donde estábamos."})

    return (
        f"Historial compactado. Resumen: {len(resumen)} chars. "
        f"Ruta preservada: {RUTA_PROYECTO[0] or 'ninguna'}"
    )


def correr_agente_claude_code(mensaje_usuario):
    global _historial
    tools_text = _formatear_herramientas_texto()
    _historial.append({"role": "user", "content": mensaje_usuario})
    mensajes = [{"role": e["role"], "content": e["content"]} for e in _historial]
    _respuesta_asistente = ""

    MAX_ITER = 10
    _reintentos_meta = 0
    for _ in range(MAX_ITER):
        # Mantener últimas 12 entradas del historial para no superar el límite de contexto
        MAX_HIST = 12
        hist = mensajes if len(mensajes) <= MAX_HIST else mensajes[-MAX_HIST:]
        partes = [SYSTEM_PROMPT, "", tools_text]
        for msg in hist:
            partes.append("")
            if msg["role"] == "user":
                partes.append(f"[USUARIO]: {msg['content']}")
            elif msg["role"] == "assistant":
                partes.append(f"[ASISTENTE]: {msg['content']}")
            elif msg["role"] == "tool_result":
                partes.append(f"[RESULTADO DE HERRAMIENTAS]:\n{msg['content']}")
        partes.append("\n[ASISTENTE]:")
        prompt = "\n".join(partes)

        # Llamar a claude -p con stdin piped para poder responder prompts de permiso desde la GUI
        cwd_proceso = RUTA_PROYECTO[0] if RUTA_PROYECTO[0] and os.path.isdir(RUTA_PROYECTO[0]) else None
        salida = ""
        try:
            # --disallowedTools bloquea herramientas de escritura nativas para forzar salida texto.
            # --append-system-prompt inyecta la instrucción como system prompt real (mayor prioridad).
            instruccion_markdown = (
                "MODO TEXTO PURO: No puedes usar las herramientas nativas Write ni Edit. "
                "Para crear o modificar archivos DEBES usar este formato exacto en tu respuesta:\n\n"
                "**`nombre.archivo`**\n```extension\ncontenido completo del archivo\n```\n\n"
                "Un bloque por archivo. NUNCA expliques permisos, NUNCA menciones la terminal, "
                "NUNCA ofrezcas opciones A/B. Solo entrega el código en bloques markdown."
            )
            cmd = [
                "claude", "-p", prompt,
                "--disallowedTools", "Write,Edit",
                "--append-system-prompt", instruccion_markdown,
            ]
            process = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE, stderr=subprocess.PIPE, stdin=subprocess.PIPE,
                text=True, cwd=cwd_proceso
            )

            _PERM_RE = re.compile(
                r"(Do you want to allow|wants? to (write|create|edit|modify|delete)"
                r"|Allow\?|\[y/n\]|\[yes/no\]|permission (required|needed|requested)"
                r"|allow .{0,60} to|Permitir|¿Permitir)",
                re.IGNORECASE
            )

            stdout_chunks = []
            stderr_buf    = []
            perm_pending  = threading.Event()

            def _leer_stdout():
                for line in iter(process.stdout.readline, ""):
                    stdout_chunks.append(line)

            def _leer_stderr():
                for line in iter(process.stderr.readline, ""):
                    stderr_buf.append(line)
                    if _PERM_RE.search(line):
                        perm_pending.set()

            t_out = threading.Thread(target=_leer_stdout, daemon=True)
            t_err = threading.Thread(target=_leer_stderr, daemon=True)
            t_out.start()
            t_err.start()

            while process.poll() is None:
                if _detener_agente.is_set():
                    process.terminate()
                    break
                if perm_pending.wait(timeout=0.2):
                    perm_pending.clear()
                    texto_perm = "".join(stderr_buf[-12:]).strip()
                    stderr_buf.clear()
                    if _solicitar_confirmacion[0]:
                        aprobado = _solicitar_confirmacion[0](
                            f"Claude Code solicita permiso:\n\n{texto_perm}"
                        )
                    else:
                        aprobado = True
                    try:
                        process.stdin.write("y\n" if aprobado else "n\n")
                        process.stdin.flush()
                    except OSError:
                        pass

            # Vaciar lo que quede pendiente tras el exit
            if perm_pending.is_set():
                perm_pending.clear()
                texto_perm = "".join(stderr_buf).strip()
                if _solicitar_confirmacion[0] and texto_perm:
                    _solicitar_confirmacion[0](f"Claude Code solicitó permiso:\n\n{texto_perm}")

            t_out.join(timeout=5)
            t_err.join(timeout=5)
            salida = "".join(stdout_chunks)

            if process.returncode != 0 and not salida.strip():
                err = "".join(stderr_buf).strip()
                print(f"Error: {err or f'Código {process.returncode}'}")
                break

        except FileNotFoundError:
            print("Error: no se encontró el comando 'claude'. ¿Está Claude Code instalado?")
            break

        salida_limpia = _FILTRO_PERMISO_RE.sub("", salida).strip()
        _respuesta_asistente = salida_limpia or salida
        mensajes.append({"role": "assistant", "content": _respuesta_asistente})

        # Extraer archivos: primero <file> blocks, si no hay, markdown code fences
        archivos = _extraer_archivos(salida)
        if not archivos:
            archivos = _extraer_archivos_markdown(salida)

        # Detectar meta-explicación ANTES de mostrar al usuario: si el modelo explicó permisos
        # en lugar de entregar archivos, reintentamos silenciosamente sin mostrar el mensaje de error.
        tool_calls = _extraer_tool_calls(salida)
        if not archivos and not tool_calls and _reintentos_meta < 2 and _META_EXPLICACION_RE.search(salida):
            _reintentos_meta += 1
            mensajes.append({
                "role": "user",
                "content": (
                    "No expliques permisos ni opciones. Entrega el código de los archivos AHORA:\n\n"
                    "**`nombre.ext`**\n```\ncontenido completo del archivo\n```\n\n"
                    "Un bloque por archivo. Sin texto adicional."
                )
            })
            continue

        # Mostrar texto visible (sin tool_call, file blocks, bloques de código ni texto de permisos)
        texto_visible = re.sub(r"<tool_call>.*?(?:</tool_call>|$)", "", salida, flags=re.DOTALL)
        texto_visible = re.sub(r'<file\s+path="[^"]*">.*?</file>', "", texto_visible, flags=re.DOTALL)
        texto_visible = re.sub(r"```[^\n]*\n.*?```", "", texto_visible, flags=re.DOTALL)
        texto_visible = _FILTRO_PERMISO_RE.sub("", texto_visible).strip()
        if texto_visible:
            print(texto_visible)

        if archivos and not RUTA_PROYECTO[0]:
            # Hay archivos pero no hay ruta — pedirla ahora en lugar de descartar silenciosamente
            if _solicitar_ruta[0]:
                _solicitar_ruta[0]()
        if archivos and RUTA_PROYECTO[0]:
            lista = "\n".join(f"  • {a['path']}" for a in archivos)
            pregunta = f"El agente quiere crear {len(archivos)} archivo(s) en:\n{RUTA_PROYECTO[0]}\n\n{lista}"
            aprobado = _solicitar_confirmacion[0](pregunta) if _solicitar_confirmacion[0] else True
            if aprobado:
                escritos = []
                for archivo in archivos:
                    ruta_abs, err = _ruta_segura(archivo["path"])
                    if err:
                        print(f"[Error: {err}]")
                        continue
                    os.makedirs(os.path.dirname(ruta_abs) or RUTA_PROYECTO[0], exist_ok=True)
                    with open(ruta_abs, "w", encoding="utf-8") as f:
                        f.write(archivo["content"])
                    escritos.append(ruta_abs)
                    print(f"[Archivo creado: {ruta_abs}]")
                print(f"\nArchivos listos en: {RUTA_PROYECTO[0]}")
                nombres = [os.path.basename(r) for r in escritos]
                mensajes.append({"role": "tool_result", "content": f"Archivos creados: {', '.join(nombres)}"})
            else:
                print("[Creación de archivos cancelada.]")
                mensajes.append({"role": "tool_result", "content": "El usuario canceló la creación de archivos."})
            break  # terminar loop tras operación de archivos — el usuario inicia cambios desde el chat

        if not tool_calls:
            break  # el modelo terminó su respuesta, sin herramientas ni archivos pendientes

        resultados = []
        for tc in tool_calls:
            nombre    = tc.get("name", "")
            inputs_tc = tc.get("inputs", {})
            print(f"\n[Herramienta: {nombre}]")
            resultado = ejecutar_herramienta(nombre, inputs_tc)
            if nombre == "leer_archivo":
                print(f"[→ {inputs_tc.get('ruta_relativa', '?')}]")
            else:
                print(f"[→ {resultado}]")
            resultados.append(f"{nombre}: {resultado}")

        if resultados:
            mensajes.append({"role": "tool_result", "content": "\n".join(resultados)})

    if _respuesta_asistente:
        _historial.append({"role": "assistant", "content": _respuesta_asistente})


def _tools_a_openai():
    """Convierte TOOLS al formato OpenAI/ZAI para function calling nativo."""
    return [
        {
            "type": "function",
            "function": {
                "name": t["name"],
                "description": t["description"],
                "parameters": t["input_schema"],
            }
        }
        for t in _tools_activos()
    ]


_ZHIPU_ACTION_SUFFIX = (
    "\n\n## REGLA CRÍTICA DE EJECUCIÓN\n"
    "Debes usar tool calls para realizar acciones. "
    "NUNCA describas lo que vas a hacer sin hacerlo — "
    "si necesitas editar un archivo, llama editar_archivo AHORA. "
    "Si necesitas leer un archivo, llama leer_archivo AHORA. "
    "Decir 'voy a hacer X' sin llamar la herramienta es un error.\n"
    "Si necesitas la ruta del proyecto, llama solicitar_ruta_proyecto AHORA — "
    "NUNCA le preguntes la ruta al usuario por texto. "
    "El selector de carpeta se abrirá automáticamente.\n"
    "REGLA DE EDICIÓN: SIEMPRE llama leer_archivo ANTES de editar_archivo. "
    "Nunca escribas texto_original de memoria — copia el texto exacto del archivo leído.\n"
    "REGLA PDF/IMPRESIÓN: Si el usuario menciona PDF, impresión, colores en PDF, "
    "header blanco, fondo que desaparece: llama diagnosticar_impresion() PRIMERO. "
    "NUNCA busques en script.js el botón de descarga — el bug está en CSS @media print, no en JS de restauración."
)


def correr_agente_zhipu(mensaje_usuario):
    global _historial
    if client is None:
        print(
            "Error: cliente zhipu no inicializado.\n"
            "  Verifica que ZAI_API_KEY esté seteada y que zai-sdk esté instalado.\n"
            "  Ejecuta: export ZAI_API_KEY=tu_clave  y  pip install zai-sdk"
        )
        return
    _historial.append({"role": "user", "content": mensaje_usuario})
    messages = [{"role": "system", "content": SYSTEM_PROMPT + _ZHIPU_ACTION_SUFFIX}]
    messages += [{"role": e["role"], "content": e["content"]} for e in _historial
                 if e["role"] in ("user", "assistant")]

    import zai as _zai
    MAX_ITER = 10
    respuesta_final = ""
    tools_openai = _tools_a_openai()
    tool_log = []
    _ya_compacto = False

    for _ in range(MAX_ITER):
        if _detener_agente.is_set():
            break
        try:
            response = client.chat.completions.create(
                model=MODELO,
                messages=messages,
                tools=tools_openai,
                tool_choice="auto",
            )
        except _zai.core.APIStatusError as e:
            print(f"Error API zhipu ({e.status_code}): {e}")
            break
        except _zai.core.APITimeoutError:
            print("Error: timeout en zhipu. Reintenta.")
            break
        except Exception as e:
            print(f"Error zhipu: {e}")
            break

        if _detener_agente.is_set():
            break

        choice = response.choices[0]
        msg = choice.message
        finish_reason = getattr(choice, "finish_reason", None)
        native_tool_calls = getattr(msg, "tool_calls", None) or []

        if not native_tool_calls:
            content = (msg.content or "").strip()
            # Quitar bloques <think>...</think> del modelo GLM (chain-of-thought interno)
            content = re.sub(r"<think>.*?</think>", "", content, flags=re.DOTALL).strip()
            if not content:
                tokens_aprox = sum(len(str(m.get("content", "") or "")) for m in messages)
                es_limite = finish_reason == "length" or tokens_aprox > 10_000
                if es_limite and not _ya_compacto:
                    _ya_compacto = True
                    print(f"\n[Contexto demasiado largo ({tokens_aprox} chars) — compactando historial automáticamente...]")
                    resultado_compact = compactar_historial()
                    print(f"[{resultado_compact}]")
                    # Volver a agregar el mensaje actual (compactar_historial lo eliminó)
                    _historial.append({"role": "user", "content": mensaje_usuario})
                    messages = [{"role": "system", "content": SYSTEM_PROMPT + _ZHIPU_ACTION_SUFFIX}]
                    messages += [
                        {"role": e["role"], "content": e["content"]}
                        for e in _historial if e["role"] in ("user", "assistant")
                    ]
                    continue
                motivo = f"finish_reason='{finish_reason}'" if finish_reason else "sin finish_reason"
                print(
                    f"\n[Zhipu devolvió respuesta vacía — {motivo}]\n"
                    f"  Causa probable: filtro de contenido, límite de tokens, o contexto demasiado largo.\n"
                    f"  Mensajes en contexto: {len(messages)} | Tokens aprox: {tokens_aprox}\n"
                    f"  Intenta reiniciar la conversación o reducir el historial."
                )
                break
            respuesta_final = content
            print("\nAgente:", respuesta_final)
            messages.append({"role": "assistant", "content": content})
            break

        # El modelo emitió tool calls nativos
        def _limpiar_nombre(n):
            n = re.split(r'[(\s<]', n or "")[0]
            return re.sub(r'[^\w]', '', n)

        messages.append({"role": "assistant", "content": msg.content or "", "tool_calls": [
            {
                "id": tc.id,
                "type": "function",
                "function": {"name": _limpiar_nombre(tc.function.name), "arguments": tc.function.arguments},
            }
            for tc in native_tool_calls
        ]})

        for tc in native_tool_calls:
            nombre_raw = tc.function.name or ""
            # GLM a veces mete args y tags XML en el name field — extraer solo el identificador
            nombre = re.split(r'[(\s<]', nombre_raw)[0].strip()
            nombre = re.sub(r'[^\w]', '', nombre)

            raw_args = (tc.function.arguments or "").strip()
            # Si arguments llega vacío pero el name tenía args embebidos, recuperarlos
            if not raw_args and "(" in nombre_raw:
                m = re.search(r'\((\{.*?\})\)', nombre_raw, re.DOTALL)
                if m:
                    raw_args = m.group(1)
            try:
                inputs_tc = json.loads(raw_args or "{}")
            except Exception:
                inputs_tc = {}
            if _detener_agente.is_set():
                break
            print(f"[Herramienta: {nombre}]")
            resultado = ejecutar_herramienta(nombre, inputs_tc)
            resultado_str = str(resultado)
            if nombre == "leer_archivo":
                print(f"[→ {inputs_tc.get('ruta_relativa', '?')}]")
                tool_log.append(f"{nombre}({inputs_tc.get('ruta_relativa', '?')}) → [{len(resultado_str)} chars leídos]")
            else:
                preview = resultado_str[:300] + ("…" if len(resultado_str) > 300 else "")
                print(f"[→ {preview}]")
                tool_log.append(f"{nombre}({json.dumps(inputs_tc, ensure_ascii=False, default=str)}) → {resultado_str[:400]}")
            messages.append({
                "role": "tool",
                "tool_call_id": tc.id,
                "content": resultado_str,
            })

    # Guardar resumen de tools al historial para que el modelo recuerde contexto en el próximo turno
    if tool_log:
        resumen_tools = "Acciones realizadas este turno:\n" + "\n".join(f"  • {t}" for t in tool_log)
        _historial.append({"role": "assistant", "content": resumen_tools})
        if not respuesta_final:
            print("\nAcciones completadas:\n" + "\n".join(f"  • {t}" for t in tool_log))
    if respuesta_final:
        _historial.append({"role": "assistant", "content": respuesta_final})


def correr_agente(mensaje_usuario):
    print(f"[Backend: {BACKEND} | {MODOS_ETIQUETAS[MODO_ACTUAL]}]\n")
    if BACKEND == "anthropic":
        correr_agente_anthropic(mensaje_usuario)
    elif BACKEND == "ollama":
        correr_agente_ollama(mensaje_usuario)
    elif BACKEND == "zhipu":
        correr_agente_zhipu(mensaje_usuario)
    elif BACKEND == "claude-code":
        correr_agente_claude_code(mensaje_usuario)


# ============================================================
# GUI
# ============================================================

if __name__ == "__main__":
    import tkinter as tk
    from tkinter import scrolledtext
    import threading
    import queue
    import builtins

    msg_queue = queue.Queue()
    _print_original = builtins.print

    # Los eventos se crean frescos por cada solicitud para evitar condiciones de carrera.

    def _mostrar_dialogo_permiso(nombre_modo, etiqueta, contenido, callback):
        dialogo = tk.Toplevel(ventana)
        dialogo.title("Permiso requerido — Nuevo agente")
        dialogo.geometry("620x480")
        dialogo.resizable(False, False)
        dialogo.configure(bg="#1e1e2e")

        tk.Label(
            dialogo, text="El agente quiere crear un nuevo experto:",
            font=("Helvetica", 13, "bold"), bg="#1e1e2e", fg="#ffffff"
        ).pack(pady=(18, 4), padx=18, anchor="w")

        tk.Label(
            dialogo, text=f"Modo: {nombre_modo}   |   Rol: {etiqueta}",
            font=("Helvetica", 12), bg="#1e1e2e", fg="#7dd3fc"
        ).pack(padx=18, anchor="w")

        tk.Label(
            dialogo, text="Contenido del prompt:",
            font=("Helvetica", 11), bg="#1e1e2e", fg="#a0a0b0"
        ).pack(pady=(14, 2), padx=18, anchor="w")

        preview = scrolledtext.ScrolledText(
            dialogo, wrap=tk.WORD, font=("Helvetica", 11), height=13,
            bg="#2a2a3e", fg="#e0e0f0", insertbackground="white"
        )
        preview.insert(tk.END, contenido)
        preview.config(state="disabled")
        preview.pack(fill=tk.BOTH, expand=True, padx=18)

        frame_btns = tk.Frame(dialogo, bg="#1e1e2e")
        frame_btns.pack(pady=14)

        def aceptar():
            callback(True)
            dialogo.destroy()

        def rechazar():
            callback(False)
            dialogo.destroy()

        dialogo.protocol("WM_DELETE_WINDOW", rechazar)

        btn_aceptar = tk.Label(
            frame_btns, text="✓  Aceptar",
            font=("Helvetica", 12, "bold"), bg="#16a34a", fg="white",
            padx=20, pady=6, cursor="hand2"
        )
        btn_aceptar.bind("<Button-1>", lambda e: aceptar())
        btn_aceptar.bind("<Enter>", lambda e: btn_aceptar.config(bg="#15803d"))
        btn_aceptar.bind("<Leave>", lambda e: btn_aceptar.config(bg="#16a34a"))
        btn_aceptar.pack(side=tk.LEFT, padx=10)

        btn_rechazar = tk.Label(
            frame_btns, text="✗  Rechazar",
            font=("Helvetica", 12, "bold"), bg="#dc2626", fg="white",
            padx=20, pady=6, cursor="hand2"
        )
        btn_rechazar.bind("<Button-1>", lambda e: rechazar())
        btn_rechazar.bind("<Enter>", lambda e: btn_rechazar.config(bg="#b91c1c"))
        btn_rechazar.bind("<Leave>", lambda e: btn_rechazar.config(bg="#dc2626"))
        btn_rechazar.pack(side=tk.LEFT, padx=10)

        dialogo.update_idletasks()
        dialogo.grab_set()
        dialogo.lift()
        dialogo.attributes("-topmost", True)
        dialogo.focus_force()

    def _solicitar_permiso_desde_hilo(nombre_modo, etiqueta, contenido):
        evento   = threading.Event()
        resultado = [False]
        def callback(val):
            resultado[0] = val
            evento.set()
        msg_queue.put(("permiso", (nombre_modo, etiqueta, contenido, callback)))
        evento.wait()
        return resultado[0]

    _solicitar_permiso[0] = _solicitar_permiso_desde_hilo

    def _mostrar_selector_ruta(callback):
        from tkinter import filedialog

        dialogo = tk.Toplevel(ventana)
        dialogo.title("Carpeta del proyecto")
        dialogo.geometry("540x240")
        dialogo.resizable(True, False)
        dialogo.configure(bg="#1e1e2e")
        dialogo.grab_set()
        dialogo.lift()
        dialogo.attributes("-topmost", True)
        dialogo.focus_force()

        tk.Label(
            dialogo, text="Ruta del proyecto:",
            font=("Helvetica", 12, "bold"), bg="#1e1e2e", fg="#e2e8f0"
        ).pack(pady=(18, 4), padx=20, anchor="w")

        tk.Label(
            dialogo,
            text="Escribe un nombre (se crea en el directorio del agente) o usa Examinar.",
            font=("Helvetica", 10), fg="#6b7280", bg="#1e1e2e", justify="left"
        ).pack(padx=20, anchor="w")

        frame_entrada = tk.Frame(dialogo, bg="#1e1e2e")
        frame_entrada.pack(fill=tk.X, padx=20, pady=(8, 0))

        entrada_ruta = tk.Entry(
            frame_entrada, font=("Helvetica", 11),
            bg="#252540", fg="#e2e8f0", insertbackground="#e2e8f0",
            relief="flat", highlightbackground="#3a3a5c",
            highlightcolor="#3b82f6", highlightthickness=1
        )
        entrada_ruta.pack(side=tk.LEFT, fill=tk.X, expand=True, ipady=6)
        entrada_ruta.focus()

        def examinar():
            dialogo.attributes("-topmost", False)
            ruta_sel = filedialog.askdirectory(
                title="Seleccionar carpeta del proyecto",
                initialdir=AGENT_DIR,
                parent=dialogo,
            )
            dialogo.attributes("-topmost", True)
            dialogo.lift()
            if ruta_sel:
                entrada_ruta.delete(0, tk.END)
                entrada_ruta.insert(0, ruta_sel)

        btn_examinar = tk.Label(
            frame_entrada, text="📂 Examinar",
            font=("Helvetica", 11, "bold"), bg="#4f46e5", fg="white",
            padx=10, pady=5, cursor="hand2"
        )
        btn_examinar.bind("<Button-1>", lambda e: examinar())
        btn_examinar.bind("<Enter>", lambda e: btn_examinar.config(bg="#4338ca"))
        btn_examinar.bind("<Leave>", lambda e: btn_examinar.config(bg="#4f46e5"))
        btn_examinar.pack(side=tk.LEFT, padx=(8, 0))

        def confirmar():
            valor = entrada_ruta.get().strip()
            if not valor:
                return
            # Ruta absoluta ingresada directamente o vía Examinar
            if os.path.isabs(valor):
                ruta = valor
            else:
                ruta = os.path.join(AGENT_DIR, valor)
            os.makedirs(ruta, exist_ok=True)
            msg_queue.put(("sistema", f"Ruta del proyecto: {ruta}\n\n"))
            dialogo.destroy()
            callback(ruta)

        def cancelar():
            dialogo.destroy()
            callback(None)

        dialogo.protocol("WM_DELETE_WINDOW", cancelar)
        entrada_ruta.bind("<Return>", lambda _: confirmar())

        frame_btns = tk.Frame(dialogo, bg="#1e1e2e")
        frame_btns.pack(pady=16)

        btn_confirmar = tk.Label(
            frame_btns, text="✓  Confirmar",
            font=("Helvetica", 12, "bold"), bg="#15803d", fg="white",
            padx=18, pady=6, cursor="hand2"
        )
        btn_confirmar.bind("<Button-1>", lambda e: confirmar())
        btn_confirmar.bind("<Enter>", lambda e: btn_confirmar.config(bg="#166534"))
        btn_confirmar.bind("<Leave>", lambda e: btn_confirmar.config(bg="#15803d"))
        btn_confirmar.pack(side=tk.LEFT, padx=10)

        btn_cancelar = tk.Label(
            frame_btns, text="✗  Cancelar",
            font=("Helvetica", 12, "bold"), bg="#b91c1c", fg="white",
            padx=18, pady=6, cursor="hand2"
        )
        btn_cancelar.bind("<Button-1>", lambda e: cancelar())
        btn_cancelar.bind("<Enter>", lambda e: btn_cancelar.config(bg="#991b1b"))
        btn_cancelar.bind("<Leave>", lambda e: btn_cancelar.config(bg="#b91c1c"))
        btn_cancelar.pack(side=tk.LEFT, padx=10)

    def _solicitar_ruta_desde_hilo():
        evento = threading.Event()
        def callback(ruta):
            RUTA_PROYECTO[0] = ruta
            evento.set()
        msg_queue.put(("ruta", callback))
        evento.wait()
        return RUTA_PROYECTO[0]

    _solicitar_ruta[0] = _solicitar_ruta_desde_hilo

    def _mostrar_dialogo_confirmacion(pregunta, callback):
        dialogo = tk.Toplevel(ventana)
        dialogo.title("El agente solicita confirmación")
        dialogo.geometry("560x320")
        dialogo.resizable(True, True)
        dialogo.configure(bg="#1e1e2e")

        tk.Label(
            dialogo, text="El agente quiere realizar una acción:",
            font=("Helvetica", 12, "bold"), bg="#1e1e2e", fg="#ffffff"
        ).pack(pady=(20, 8), padx=20, anchor="w")

        txt = scrolledtext.ScrolledText(
            dialogo, wrap=tk.WORD, font=("Helvetica", 12),
            height=8, bg="#2a2a3e", fg="#e0e0f0", relief="flat",
            insertbackground="white"
        )
        txt.insert(tk.END, pregunta)
        txt.config(state="disabled")
        txt.pack(fill=tk.BOTH, expand=True, padx=20, pady=(0, 6))

        frame_btns = tk.Frame(dialogo, bg="#1e1e2e")
        frame_btns.pack(pady=14)

        def aceptar():
            callback(True)
            dialogo.destroy()

        def rechazar():
            callback(False)
            dialogo.destroy()

        dialogo.protocol("WM_DELETE_WINDOW", rechazar)

        btn_si = tk.Label(
            frame_btns, text="✓  Sí, procede",
            font=("Helvetica", 12, "bold"), bg="#16a34a", fg="white",
            padx=16, pady=6, cursor="hand2"
        )
        btn_si.bind("<Button-1>", lambda e: aceptar())
        btn_si.bind("<Enter>", lambda e: btn_si.config(bg="#15803d"))
        btn_si.bind("<Leave>", lambda e: btn_si.config(bg="#16a34a"))
        btn_si.pack(side=tk.LEFT, padx=10)

        btn_no = tk.Label(
            frame_btns, text="✗  No, cancela",
            font=("Helvetica", 12, "bold"), bg="#dc2626", fg="white",
            padx=16, pady=6, cursor="hand2"
        )
        btn_no.bind("<Button-1>", lambda e: rechazar())
        btn_no.bind("<Enter>", lambda e: btn_no.config(bg="#b91c1c"))
        btn_no.bind("<Leave>", lambda e: btn_no.config(bg="#dc2626"))
        btn_no.pack(side=tk.LEFT, padx=10)

        dialogo.update_idletasks()
        dialogo.grab_set()
        dialogo.lift()
        dialogo.attributes("-topmost", True)
        dialogo.focus_force()

    def _solicitar_confirmacion_desde_hilo(pregunta):
        evento    = threading.Event()
        resultado = [False]
        def callback(val):
            resultado[0] = val
            evento.set()
        msg_queue.put(("confirmacion", (pregunta, callback)))
        evento.wait()
        return resultado[0]

    _solicitar_confirmacion[0] = _solicitar_confirmacion_desde_hilo

    # --- UI helpers ---
    # Referencias para sincronizar widgets definidos después de check_queue
    _sync_modo_ref        = [None]   # fn(modo_key) → actualiza dropdown de agente
    _sync_menu_modos_ref  = [None]   # fn() → reconstruye opciones del dropdown
    _refrescar_pills_ref  = [None]   # fn() → añade pills de nuevas herramientas

    def titulo_ventana():
        return f"Agente [{BACKEND}] — {MODOS_ETIQUETAS[MODO_ACTUAL]}"

    def agregar_texto(texto, tag):
        area.config(state="normal")
        area.insert(tk.END, texto, tag)
        area.see(tk.END)
        area.config(state="disabled")

    def check_queue():
        try:
            while True:
                tag, contenido = msg_queue.get_nowait()
                try:
                    if tag == "permiso":
                        nombre_modo, etiqueta, prompt_txt, cb = contenido
                        _mostrar_dialogo_permiso(nombre_modo, etiqueta, prompt_txt, cb)
                    elif tag == "ruta":
                        _mostrar_selector_ruta(contenido)
                    elif tag == "confirmacion":
                        pregunta, cb = contenido
                        _mostrar_dialogo_confirmacion(pregunta, cb)
                    elif tag == "modo":
                        ventana.title(titulo_ventana())
                        agregar_texto(f"\n— Modo: {MODOS_ETIQUETAS[contenido]} —\n\n", "modo")
                        if _sync_modo_ref[0]:
                            _sync_modo_ref[0](contenido)
                        if _sync_menu_modos_ref[0]:
                            _sync_menu_modos_ref[0]()
                    else:
                        agregar_texto(contenido, tag)
                except Exception as _ex:
                    _print_original(f"[check_queue handler error] {_ex}")
        except queue.Empty:
            pass
        try:
            if _refrescar_pills_ref[0]:
                _refrescar_pills_ref[0]()
        except Exception as _ex:
            _print_original(f"[refrescar_pills error] {_ex}")
        ventana.after(100, check_queue)

    # --- Indicador de procesamiento ---
    _procesando   = [False]
    _anim_job     = [None]
    _anim_idx     = [0]
    _anim_frames  = ["Procesando ·  ", "Procesando ·· ", "Procesando ···"]

    def _animar():
        if not _procesando[0]:
            return
        lbl_estado.config(text=_anim_frames[_anim_idx[0] % 3], fg="#cc6600")
        _anim_idx[0] += 1
        _anim_job[0] = ventana.after(400, _animar)

    def iniciar_procesando():
        _procesando[0] = True
        _anim_idx[0]   = 0
        _animar()

    def detener_procesando():
        _procesando[0] = False
        if _anim_job[0]:
            ventana.after_cancel(_anim_job[0])
            _anim_job[0] = None
        lbl_estado.config(text="")

    def enviar():
        mensaje = entrada.get("1.0", "end-1c").strip()
        if not mensaje:
            return
        entrada.delete("1.0", tk.END)

        if mensaje.startswith("/modo "):
            nombre_modo = mensaje[6:].strip()
            resultado = ejecutar_herramienta("cambiar_modo", {"modo": nombre_modo})
            ventana.title(titulo_ventana())
            agregar_texto(f"Sistema: {resultado}\n\n", "sistema")
            return

        if mensaje in ("/compact", "/compactar"):
            resultado = compactar_historial()
            agregar_texto(f"Sistema: {resultado}\n\n", "sistema")
            return

        _actualizar_btn(False)
        agregar_texto(f"Tú: {mensaje}\n", "usuario")
        iniciar_procesando()
        _detener_agente.clear()

        _n_msgs = [0]

        def print_a_cola(*args):
            texto = " ".join(str(a) for a in args)
            if texto.startswith("__MODO__"):
                msg_queue.put(("modo", texto[8:]))
            else:
                msg_queue.put(("agente", texto + "\n"))
                _n_msgs[0] += 1

        def ejecutar():
            _n_msgs[0] = 0
            builtins.print = print_a_cola
            try:
                correr_agente(mensaje)
                if _n_msgs[0] == 0:
                    msg_queue.put(("sistema", "Agente finalizó sin respuesta visible. Sin cambios realizados.\n\n"))
                else:
                    msg_queue.put(("agente", "\n"))
            except Exception as ex:
                msg_queue.put(("error", f"Error: {ex}\n"))
            finally:
                builtins.print = _print_original
                ventana.after(0, detener_procesando)
                ventana.after(0, lambda: _actualizar_btn(True))

        threading.Thread(target=ejecutar, daemon=True).start()

    # === Tema de colores (Catppuccin Mocha — contraste verificado) ===
    FONDO    = "#1e1e2e"   # base
    FONDO2   = "#313244"   # surface — claramente más claro que FONDO
    CHAT_BG  = "#181825"   # mantle — ligeramente distinto del base
    TEXTO    = "#cdd6f4"   # text — blanco-lavanda
    SUBTEXTO = "#bac2de"   # subtext — etiquetas secundarias, contraste OK
    AZUL     = "#3b82f6"   # acento primario
    AMARILLO = "#f9e2af"   # peach suave
    BORDE    = "#45475a"   # surface2 — bordes visibles sin ser duros

    # === Ventana principal ===
    ventana = tk.Tk()
    ventana.title(titulo_ventana())
    ventana.geometry("800x620")
    ventana.minsize(540, 400)
    ventana.configure(bg=FONDO)

    # Abre sobre VS Code: -topmost activado brevemente al iniciar
    ventana.attributes("-topmost", True)
    ventana.lift()
    ventana.focus_force()
    ventana.after(600, lambda: ventana.attributes("-topmost", False))

    # === Barra superior (header) ===
    header = tk.Frame(ventana, bg=FONDO2, height=48)
    header.pack(fill=tk.X, side=tk.TOP)
    header.pack_propagate(False)

    tk.Label(
        header, text="◆  Multi-Agente IA",
        font=("Helvetica", 13, "bold"), bg=FONDO2, fg=TEXTO
    ).pack(side=tk.LEFT, padx=18, pady=12)

    # Separador vertical decorativo
    tk.Frame(header, bg=BORDE, width=1).pack(side=tk.LEFT, fill=tk.Y, pady=10)

    var_backend = tk.StringVar(value=BACKEND)

    def _on_cambio_backend(nuevo_backend):
        err = inicializar_cliente(nuevo_backend)
        if err:
            agregar_texto(f"Sistema: {err}\n\n", "error")
            var_backend.set(BACKEND)
        else:
            ventana.title(titulo_ventana())
            agregar_texto(f"Sistema: Backend → {nuevo_backend} | Modelo: {MODELO}\n\n", "sistema")

    tk.Label(
        header, text="Backend:", font=("Helvetica", 10),
        bg=FONDO2, fg=SUBTEXTO
    ).pack(side=tk.RIGHT, padx=(0, 6))

    menu_backend = tk.OptionMenu(
        header, var_backend, *BACKENDS_DISPONIBLES, command=_on_cambio_backend
    )
    menu_backend.config(
        font=("Helvetica", 11), bg=FONDO2, fg=TEXTO,
        activebackground=AZUL, activeforeground="white",
        relief="flat", borderwidth=0, highlightthickness=0, padx=8
    )
    menu_backend["menu"].config(
        bg=FONDO2, fg=TEXTO, font=("Helvetica", 11),
        activebackground=AZUL, activeforeground="white"
    )
    menu_backend.pack(side=tk.RIGHT, padx=(0, 14), pady=8)

    # === Sub-barra: selector de agente + herramientas ===
    tk.Frame(ventana, bg=BORDE, height=1).pack(fill=tk.X, side=tk.TOP)

    subheader = tk.Frame(ventana, bg=FONDO, height=42)
    subheader.pack(fill=tk.X, side=tk.TOP)
    subheader.pack_propagate(False)

    # --- Selector de agente ---
    tk.Label(
        subheader, text="Agente:", font=("Helvetica", 10),
        bg=FONDO, fg=SUBTEXTO
    ).pack(side=tk.LEFT, padx=(14, 4))

    etiqueta_a_modo = {v: k for k, v in MODOS_ETIQUETAS.items()}
    var_modo = tk.StringVar(value=MODOS_ETIQUETAS[MODO_ACTUAL])
    _sync_modo_ref[0] = lambda mk: var_modo.set(MODOS_ETIQUETAS.get(mk, mk))

    def _on_cambio_agente(etiqueta):
        modo = etiqueta_a_modo.get(etiqueta)
        if not modo:
            return
        resultado = ejecutar_herramienta("cambiar_modo", {"modo": modo})
        ventana.title(titulo_ventana())
        agregar_texto(f"Sistema: {resultado}\n\n", "sistema")

    menu_agente = tk.OptionMenu(
        subheader, var_modo, *MODOS_ETIQUETAS.values(), command=_on_cambio_agente
    )
    menu_agente.config(
        font=("Helvetica", 11), bg=FONDO2, fg=TEXTO,
        activebackground=AZUL, activeforeground="white",
        relief="flat", borderwidth=0, highlightthickness=0, padx=8
    )
    menu_agente["menu"].config(
        bg=FONDO2, fg=TEXTO, font=("Helvetica", 11),
        activebackground=AZUL, activeforeground="white"
    )
    menu_agente.pack(side=tk.LEFT, pady=6)

    def _sync_menu_modos():
        etiqueta_a_modo.clear()
        etiqueta_a_modo.update({v: k for k, v in MODOS_ETIQUETAS.items()})
        menu = menu_agente["menu"]
        menu.delete(0, "end")
        for etq in MODOS_ETIQUETAS.values():
            menu.add_command(
                label=etq,
                command=tk._setit(var_modo, etq, _on_cambio_agente)
            )
    _sync_menu_modos_ref[0] = _sync_menu_modos

    def reiniciar_agente():
        global _historial
        _detener_agente.set()
        _historial.clear()
        RUTA_PROYECTO[0] = None
        area.config(state="normal")
        area.delete("1.0", tk.END)
        area.config(state="disabled")
        agregar_texto("Sistema: Agente reiniciado.\n\n", "sistema")

    tk.Button(
        subheader, text="↺ Reiniciar",
        font=("Helvetica", 11, "bold"),
        bg="#1e3a5f", fg="#7dd3fc",
        activebackground="#2a4f7a", activeforeground="#bae6fd",
        relief="flat", borderwidth=0, padx=12, pady=4,
        cursor="hand2", command=reiniciar_agente
    ).pack(side=tk.RIGHT, padx=(0, 14), pady=6)

    # Separador vertical
    tk.Frame(subheader, bg=BORDE, width=1).pack(side=tk.LEFT, fill=tk.Y, padx=12, pady=8)

    # --- Dropdown de herramientas con checkboxes ---
    tk.Label(
        subheader, text="Herramientas:", font=("Helvetica", 10),
        bg=FONDO, fg=SUBTEXTO
    ).pack(side=tk.LEFT, padx=(0, 4))

    _SYSTEM_TOOLS = {
        "cambiar_modo", "crear_agente", "crear_herramienta",
        "pedir_confirmacion", "solicitar_ruta_proyecto"
    }
    _tools_vars: dict = {}   # nombre → BooleanVar
    _tools_popup = [None]

    def _tools_btn_text():
        if not _tools_vars:
            return "sin herramientas ▾"
        n_act = sum(1 for v in _tools_vars.values() if v.get())
        return f"{n_act}/{len(_tools_vars)} activas ▾"

    tools_btn = tk.Label(
        subheader, text="sin herramientas ▾",
        font=("Helvetica", 10), bg=FONDO2, fg=TEXTO,
        padx=10, pady=4, cursor="hand2"
    )
    tools_btn.pack(side=tk.LEFT, pady=8)

    def _cerrar_tools_popup(_e=None):
        if _tools_popup[0]:
            _tools_popup[0].destroy()
            _tools_popup[0] = None

    def _on_tool_toggle(nombre):
        if _tools_vars[nombre].get():
            TOOLS_DESACTIVADAS.discard(nombre)
        else:
            TOOLS_DESACTIVADAS.add(nombre)
        tools_btn.config(text=_tools_btn_text())

    def _abrir_tools_popup(_e=None):
        if _tools_popup[0]:
            _cerrar_tools_popup()
            return
        if not _tools_vars:
            return

        popup = tk.Toplevel(ventana)
        popup.overrideredirect(True)
        popup.attributes("-topmost", True)

        # Marco con borde
        border = tk.Frame(popup, bg=BORDE, padx=1, pady=1)
        border.pack(fill=tk.BOTH, expand=True)
        inner = tk.Frame(border, bg=FONDO2)
        inner.pack(fill=tk.BOTH, expand=True)

        for nombre, var in _tools_vars.items():
            cb = tk.Checkbutton(
                inner, text=nombre, variable=var,
                bg=FONDO2, fg=TEXTO,
                selectcolor=FONDO,
                activebackground=FONDO2, activeforeground=TEXTO,
                font=("Helvetica", 11), anchor="w",
                command=lambda n=nombre: _on_tool_toggle(n)
            )
            cb.pack(fill=tk.X, padx=10, pady=4)

        popup.update_idletasks()
        x = tools_btn.winfo_rootx()
        y = tools_btn.winfo_rooty() + tools_btn.winfo_height() + 2
        popup.geometry(f"+{x}+{y}")
        _tools_popup[0] = popup

    def _check_click_fuera(e):
        if not _tools_popup[0]:
            return
        if e.widget is tools_btn:
            return  # el botón maneja su propio toggle
        px, py = _tools_popup[0].winfo_rootx(), _tools_popup[0].winfo_rooty()
        _tools_popup[0].update_idletasks()
        pw, ph = _tools_popup[0].winfo_width(), _tools_popup[0].winfo_height()
        if not (px <= e.x_root <= px + pw and py <= e.y_root <= py + ph):
            _cerrar_tools_popup()

    ventana.bind("<Button-1>", _check_click_fuera, add="+")
    tools_btn.bind("<Button-1>", _abrir_tools_popup)

    def _refrescar_pills():
        nuevos = [t for t in TOOLS
                  if t["name"] not in _SYSTEM_TOOLS and t["name"] not in _tools_vars]
        for tool in nuevos:
            _tools_vars[tool["name"]] = tk.BooleanVar(value=True)
        if nuevos:
            tools_btn.config(text=_tools_btn_text())
            if _tools_popup[0]:
                _cerrar_tools_popup()  # se reabrirá con las nuevas herramientas

    _refrescar_pills()
    _refrescar_pills_ref[0] = _refrescar_pills

    tk.Frame(ventana, bg=BORDE, height=1).pack(fill=tk.X, side=tk.TOP)

    # === Barra de estado (procesando) — debe empaquetarse ANTES que area ===
    lbl_estado = tk.Label(
        ventana, text="", font=("Helvetica", 10, "italic"),
        anchor="w", bg=FONDO2, fg=AMARILLO, padx=18, pady=5
    )
    lbl_estado.pack(fill=tk.X, side=tk.BOTTOM)

    # === Barra de entrada — debe empaquetarse ANTES que area ===
    frame = tk.Frame(ventana, bg=FONDO)
    frame.pack(fill=tk.X, side=tk.BOTTOM, padx=14, pady=10)

    _btn_activo = [True]
    _btn_detener_ref = [None]

    def _actualizar_btn(activo):
        _btn_activo[0] = activo
        if activo:
            btn_enviar.config(bg="#1d4ed8", fg="#ffffff", cursor="hand2")
            if _btn_detener_ref[0]:
                _btn_detener_ref[0].config(text="◼ Detener", bg=BORDE, fg=SUBTEXTO, cursor="arrow")
        else:
            btn_enviar.config(bg=BORDE, fg=SUBTEXTO, cursor="arrow")
            if _btn_detener_ref[0]:
                _btn_detener_ref[0].config(text="◼ Detener", bg="#dc2626", fg="#ffffff", cursor="hand2")

    btn_enviar = tk.Label(
        frame, text="Enviar",
        font=("Helvetica", 13, "bold"),
        bg="#1d4ed8", fg="#ffffff",
        padx=26
    )
    btn_enviar.bind("<Button-1>", lambda e: enviar() if _btn_activo[0] else None)
    btn_enviar.bind("<Enter>", lambda e: btn_enviar.config(bg="#1e40af") if _btn_activo[0] else None)
    btn_enviar.bind("<Leave>", lambda e: btn_enviar.config(bg="#1d4ed8") if _btn_activo[0] else None)
    btn_enviar.pack(side=tk.RIGHT, fill=tk.Y, pady=8)

    def _detener():
        if _btn_activo[0]:
            return  # agente no está corriendo
        _detener_agente.set()
        btn_detener.config(text="⏳ Deteniendo...", bg="#92400e", cursor="arrow")
        msg_queue.put(("sistema", "[Detención solicitada — esperando respuesta API...]\n"))

    btn_detener = tk.Label(
        frame, text="◼ Detener",
        font=("Helvetica", 13, "bold"),
        bg=BORDE, fg=SUBTEXTO,
        padx=14
    )
    btn_detener.bind("<Button-1>", lambda e: _detener() if not _btn_activo[0] else None)
    btn_detener.bind("<Enter>", lambda e: btn_detener.config(bg="#b91c1c") if not _btn_activo[0] else None)
    btn_detener.bind("<Leave>", lambda e: btn_detener.config(bg="#dc2626") if not _btn_activo[0] else None)
    # Nota: _btn_activo[0]=True significa enviar activo (agente en reposo).
    #       _btn_activo[0]=False significa agente procesando → detener habilitado.
    btn_detener.pack(side=tk.RIGHT, fill=tk.Y, pady=8, padx=(0, 6))
    _btn_detener_ref[0] = btn_detener

    entrada_frame = tk.Frame(
        frame, bg=FONDO2,
        highlightbackground=BORDE, highlightcolor=AZUL, highlightthickness=1
    )
    entrada_frame.pack(side=tk.LEFT, fill=tk.BOTH, expand=True, padx=(0, 12), pady=8)

    lbl_hint = tk.Label(
        entrada_frame,
        text="Shift+Enter: nueva línea",
        font=("Helvetica", 8), bg=FONDO2, fg="#6b7280",
        anchor="e", padx=6
    )
    lbl_hint.pack(side=tk.BOTTOM, fill=tk.X, pady=(0, 2))

    entrada_scroll = tk.Scrollbar(entrada_frame, orient=tk.VERTICAL)
    entrada_scroll.pack(side=tk.RIGHT, fill=tk.Y)

    entrada = tk.Text(
        entrada_frame, font=("Helvetica", 14),
        bg=FONDO2, fg=TEXTO, insertbackground=TEXTO,
        relief="flat", wrap=tk.WORD, height=3,
        yscrollcommand=entrada_scroll.set,
        highlightthickness=0, borderwidth=0,
        padx=8, pady=6,
    )
    entrada.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)
    entrada_scroll.config(command=entrada.yview)

    def _on_return(event):
        enviar()
        return "break"

    def _on_shift_return(event):
        entrada.insert(tk.INSERT, "\n")
        return "break"

    entrada.bind("<Shift-Return>", _on_shift_return)
    entrada.bind("<Return>", _on_return)
    entrada.focus()

    # === Área de chat — última en empaquetar para respetar los side=BOTTOM ===
    area = scrolledtext.ScrolledText(
        ventana, state="disabled", wrap=tk.WORD,
        font=("Menlo", 12), bg=CHAT_BG, fg=TEXTO,
        insertbackground=TEXTO, selectbackground=AZUL,
        relief="flat", padx=18, pady=14,
        spacing1=1, spacing3=5, borderwidth=0
    )
    area.pack(fill=tk.BOTH, expand=True)

    area.tag_config("usuario", foreground="#89b4fa", font=("Menlo", 12, "bold"))
    area.tag_config("agente",  foreground="#a6e3a1", font=("Menlo", 12))
    area.tag_config("error",   foreground="#f38ba8", font=("Menlo", 12))
    area.tag_config("modo",    foreground="#cba6f7", font=("Menlo", 12, "italic"))
    area.tag_config("sistema", foreground="#a6adc8", font=("Menlo", 11, "italic"))

    check_queue()
    ventana.mainloop()
